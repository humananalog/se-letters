"""Document processing service for the SE Letters project."""

import time
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pytesseract
from pdf2image import convert_from_path
import docx2txt

from ..core.config import Config
from ..core.exceptions import FileProcessingError
from ..models.document import Document
from ..utils.logger import get_logger

# Enhanced image processor has been archived
IMAGE_PROCESSOR_AVAILABLE = False

logger = get_logger(__name__)


class DocumentProcessor:
    """Enhanced document processor with comprehensive fallback mechanisms and embedded image processing."""

    def __init__(self, config: Config) -> None:
        """Initialize the document processor.
        
        Args:
            config: Configuration instance.
        """
        self.config = config
        self.supported_formats = config.data.supported_formats
        
        # Enhanced image processor has been archived
        self.image_processor = None
        logger.info("Using standard document processing without enhanced image extraction")

    def process_document(self, file_path: Path) -> Optional[Document]:
        """Process a single document and extract text with robust fallback and embedded image processing.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Document instance if successful, None otherwise.
        """
        start_time = time.time()
        
        try:
            logger.info(f"Processing document: {file_path}")
            
            # Check if file exists and is supported
            if not file_path.exists():
                raise FileProcessingError(f"File not found: {file_path}")
            
            if file_path.suffix.lower() not in self.supported_formats:
                raise FileProcessingError(
                    f"Unsupported format: {file_path.suffix}"
                )
            
            # Extract text using robust methods
            result = self._extract_text_robust(file_path)
            
            if not result["success"]:
                logger.warning(
                    f"Text extraction failed for {file_path}: "
                    f"{result['error']}"
                )
                return None
            
            # Enhanced image processing has been archived
            if result["success"]:
                logger.debug("Enhanced image processing disabled")
            
            processing_time = time.time() - start_time
            
            # Enhanced image processing has been archived
            # No image processing results to add to metadata
            
            # Create document
            document = Document.from_file(
                file_path=file_path,
                text=result["text"],
                processing_time=processing_time,
                metadata=result["metadata"],
            )
            
            logger.info(
                f"Successfully processed {file_path} "
                f"({len(result['text'])} chars, {processing_time:.2f}s) "
                f"using {result['method_used']}"
            )
            
            return document
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(
                f"Failed to process {file_path} "
                f"after {processing_time:.2f}s: {e}"
            )
            return None

    def _enhance_text_with_image_content(self, text_result: Dict[str, Any], image_results: Dict[str, Any]) -> Dict[str, Any]:
        """Enhance extracted text with content from embedded images.
        
        Args:
            text_result: Original text extraction result.
            image_results: Image processing results.
            
        Returns:
            Enhanced text result.
        """
        if not image_results["modernization_content"]:
            return text_result
        
        # Create modernization summary
        modernization_summary = self.image_processor.create_modernization_summary(image_results)
        
        # Append image-extracted content to text
        enhanced_text = text_result["text"]
        enhanced_text += "\n\n[EMBEDDED IMAGE CONTENT - MODERNIZATION PATHS]\n"
        enhanced_text += "=" * 50 + "\n"
        
        # Add product mappings
        if modernization_summary["all_product_mappings"]:
            enhanced_text += "\nPRODUCT REPLACEMENT MAPPINGS:\n"
            for mapping in modernization_summary["all_product_mappings"]:
                enhanced_text += f"• {mapping['obsolete_code']} → {mapping['replacement_code']} ({mapping['mapping_type']})\n"
        
        # Add replacement tables
        if modernization_summary["all_replacement_tables"]:
            enhanced_text += "\nREPLACEMENT TABLES:\n"
            for i, table in enumerate(modernization_summary["all_replacement_tables"]):
                enhanced_text += f"\nTable {i+1} ({table['type']}):\n"
                enhanced_text += f"Headers: {', '.join(table['headers'])}\n"
                enhanced_text += f"Rows: {table['row_count']}\n"
                
                # Add first few rows as examples
                for j, row in enumerate(table['rows'][:3]):  # First 3 rows
                    row_text = " | ".join([f"{k}: {v}" for k, v in row.items()])
                    enhanced_text += f"  Row {j+1}: {row_text}\n"
                
                if len(table['rows']) > 3:
                    enhanced_text += f"  ... and {len(table['rows']) - 3} more rows\n"
        
        # Add modernization paths
        if modernization_summary["all_modernization_paths"]:
            enhanced_text += "\nMODERNIZATION PATHS:\n"
            for path in modernization_summary["all_modernization_paths"]:
                enhanced_text += f"• {path['description']}\n"
        
        # Add extraction quality info
        enhanced_text += f"\nImage Extraction Quality: {modernization_summary['extraction_quality'].upper()}\n"
        enhanced_text += f"Confidence Score: {modernization_summary['confidence_score']:.1f}%\n"
        enhanced_text += f"Images with Modernization Content: {modernization_summary['images_with_modernization_content']}\n"
        
        # Update result
        text_result["text"] = enhanced_text
        text_result["metadata"]["image_enhanced"] = True
        text_result["metadata"]["modernization_summary"] = modernization_summary
        
        logger.info(
            f"Enhanced text with {len(modernization_summary['all_product_mappings'])} "
            f"product mappings and {len(modernization_summary['all_replacement_tables'])} "
            f"tables from embedded images"
        )
        
        return text_result

    def _extract_text_robust(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using comprehensive fallback mechanisms.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Dictionary with extraction results.
        """
        suffix = file_path.suffix.lower()
        base_metadata = {
            "format": suffix,
            "file_size": file_path.stat().st_size,
            "methods_attempted": [],
            "method_used": "none",
            "fallback_used": False
        }
        
        if suffix == ".pdf":
            return self._extract_pdf_robust(file_path, base_metadata)
        elif suffix == ".docx":
            return self._extract_docx_robust(file_path, base_metadata)
        elif suffix == ".doc":
            return self._extract_doc_robust(file_path, base_metadata)
        else:
            return {
                "success": False,
                "text": "",
                "metadata": base_metadata,
                "method_used": "none",
                "error": f"Unsupported format: {suffix}"
            }

    def _extract_pdf_robust(self, file_path: Path, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from PDF with multiple fallback methods.
        
        Args:
            file_path: Path to PDF file.
            metadata: Base metadata dictionary.
            
        Returns:
            Dictionary with extraction results.
        """
        methods = [
            ("PyMuPDF", self._extract_pdf_pymupdf),
            ("pdfplumber", self._extract_pdf_pdfplumber),
            ("PyPDF2", self._extract_pdf_pypdf2),
            ("OCR", self._extract_pdf_ocr),
            ("Fallback", self._create_pdf_fallback)
        ]
        
        for method_name, method_func in methods:
            try:
                metadata["methods_attempted"].append(method_name)
                text, method_metadata = method_func(file_path)
                
                if text and text.strip():
                    metadata.update(method_metadata)
                    metadata["method_used"] = method_name
                    metadata["fallback_used"] = method_name in ["OCR", "Fallback"]
                    
                    return {
                        "success": True,
                        "text": text,
                        "metadata": metadata,
                        "method_used": method_name,
                        "error": None
                    }
                    
            except Exception as e:
                logger.debug(f"PDF method {method_name} failed for {file_path}: {e}")
                continue
        
        return {
            "success": False,
            "text": "",
            "metadata": metadata,
            "method_used": "none",
            "error": "All PDF extraction methods failed"
        }

    def _extract_docx_robust(self, file_path: Path, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from DOCX with multiple fallback methods.
        
        Args:
            file_path: Path to DOCX file.
            metadata: Base metadata dictionary.
            
        Returns:
            Dictionary with extraction results.
        """
        methods = [
            ("python-docx", self._extract_docx_python_docx),
            ("docx2txt", self._extract_docx_docx2txt),
            ("LibreOffice", self._extract_docx_libreoffice),
            ("Fallback", self._create_docx_fallback)
        ]
        
        for method_name, method_func in methods:
            try:
                metadata["methods_attempted"].append(method_name)
                text, method_metadata = method_func(file_path)
                
                if text and text.strip():
                    metadata.update(method_metadata)
                    metadata["method_used"] = method_name
                    metadata["fallback_used"] = method_name in ["Fallback"]
                    
                    return {
                        "success": True,
                        "text": text,
                        "metadata": metadata,
                        "method_used": method_name,
                        "error": None
                    }
                    
            except Exception as e:
                logger.debug(f"DOCX method {method_name} failed for {file_path}: {e}")
                continue
        
        return {
            "success": False,
            "text": "",
            "metadata": metadata,
            "method_used": "none",
            "error": "All DOCX extraction methods failed"
        }

    def _extract_doc_robust(self, file_path: Path, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from DOC with comprehensive fallback methods.
        
        Args:
            file_path: Path to DOC file.
            metadata: Base metadata dictionary.
            
        Returns:
            Dictionary with extraction results.
        """
        methods = [
            ("LibreOffice", self._extract_doc_libreoffice),
            ("python-docx", self._extract_doc_python_docx),
            ("docx2txt", self._extract_doc_docx2txt),
            ("OCR", self._extract_doc_ocr),
            ("Fallback", self._create_doc_fallback)
        ]
        
        for method_name, method_func in methods:
            try:
                metadata["methods_attempted"].append(method_name)
                text, method_metadata = method_func(file_path)
                
                if text and text.strip():
                    metadata.update(method_metadata)
                    metadata["method_used"] = method_name
                    metadata["fallback_used"] = method_name in ["OCR", "Fallback"]
                    
                    return {
                        "success": True,
                        "text": text,
                        "metadata": metadata,
                        "method_used": method_name,
                        "error": None
                    }
                    
            except Exception as e:
                logger.debug(f"DOC method {method_name} failed for {file_path}: {e}")
                continue
        
        return {
            "success": False,
            "text": "",
            "metadata": metadata,
            "method_used": "none",
            "error": "All DOC extraction methods failed"
        }

    # PDF extraction methods
    def _extract_pdf_pymupdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyMuPDF."""
        import fitz
        
        doc = fitz.open(file_path)
        text_parts = []
        metadata = {"pages": len(doc)}
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        doc.close()
        return "\n\n".join(text_parts), metadata

    def _extract_pdf_pdfplumber(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using pdfplumber."""
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            metadata = {"pages": len(pdf.pages)}
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts), metadata

    def _extract_pdf_pypdf2(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyPDF2."""
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {"pages": len(pdf_reader.pages)}
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts), metadata
            
        except ImportError:
            raise Exception("PyPDF2 not installed")

    def _extract_pdf_ocr(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using OCR."""
        images = convert_from_path(file_path)
        text_parts = []
        metadata = {"pages": len(images), "ocr_used": True}
        
        for i, image in enumerate(images):
            try:
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(text)
            except Exception as e:
                logger.debug(f"OCR failed for page {i}: {e}")
        
        return "\n\n".join(text_parts), metadata

    def _create_pdf_fallback(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Create fallback content for PDF files."""
        filename = file_path.name
        content = f"[PDF Document: {filename}]\n"
        content += f"File size: {file_path.stat().st_size} bytes\n"
        content += self._infer_content_from_filename(filename)
        
        return content, {"fallback_content": True}

    # DOCX extraction methods
    def _extract_docx_python_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using python-docx."""
        doc = DocxDocument(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # If no content found, create minimal content
        if not text_parts:
            text_parts = [f"[DOCX Document: {file_path.name}]"]
            text_parts.append(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
            text_parts.append("Content appears to be primarily formatting or images.")
        
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }
        
        return "\n\n".join(text_parts), metadata

    def _extract_docx_docx2txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using docx2txt."""
        text = docx2txt.process(str(file_path))
        return text, {"extraction_library": "docx2txt"}

    def _extract_docx_libreoffice(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using LibreOffice."""
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            result = subprocess.run([
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless",
                "--convert-to", "txt",
                "--outdir", str(temp_path),
                str(file_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
            txt_files = list(temp_path.glob("*.txt"))
            if not txt_files:
                raise Exception("No text file produced")
            
            text = txt_files[0].read_text(encoding='utf-8')
            return text, {"converted_via": "libreoffice"}

    def _create_docx_fallback(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Create fallback content for DOCX files."""
        filename = file_path.name
        content = f"[DOCX Document: {filename}]\n"
        content += f"File size: {file_path.stat().st_size} bytes\n"
        content += self._infer_content_from_filename(filename)
        
        return content, {"fallback_content": True}

    # DOC extraction methods
    def _extract_doc_libreoffice(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOC using LibreOffice."""
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Convert DOC to DOCX first
            result = subprocess.run([
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(temp_path),
                str(file_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
            docx_files = list(temp_path.glob("*.docx"))
            if not docx_files:
                raise Exception("No DOCX file produced")
            
            # Extract text from converted DOCX
            text, metadata = self._extract_docx_python_docx(docx_files[0])
            metadata["converted_from_doc"] = True
            
            return text, metadata

    def _extract_doc_python_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOC using python-docx (may work for some DOC files)."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return "\n\n".join(text_parts), {"extraction_method": "docx_on_doc"}
            
        except Exception as e:
            raise Exception(f"python-docx failed on DOC: {e}")

    def _extract_doc_docx2txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOC using docx2txt."""
        try:
            text = docx2txt.process(str(file_path))
            return text, {"extraction_library": "docx2txt_on_doc"}
        except Exception as e:
            raise Exception(f"docx2txt failed on DOC: {e}")

    def _extract_doc_ocr(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOC using OCR (convert to images first)."""
        # This is a last resort - convert DOC to PDF then to images
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Convert DOC to PDF
            result = subprocess.run([
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(temp_path),
                str(file_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception("Could not convert DOC to PDF for OCR")
            
            pdf_files = list(temp_path.glob("*.pdf"))
            if not pdf_files:
                raise Exception("No PDF file produced for OCR")
            
            # OCR the PDF
            text, metadata = self._extract_pdf_ocr(pdf_files[0])
            metadata["ocr_via_pdf_conversion"] = True
            
            return text, metadata

    def _create_doc_fallback(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Create fallback content for DOC files."""
        filename = file_path.name
        content = f"[DOC Document: {filename}]\n"
        content += f"File size: {file_path.stat().st_size} bytes\n"
        content += self._infer_content_from_filename(filename)
        
        return content, {"fallback_content": True}

    def _infer_content_from_filename(self, filename: str) -> str:
        """Infer content type and context from filename."""
        filename_upper = filename.upper()
        content_lines = []
        
        # Product range detection
        product_ranges = []
        range_patterns = {
            "PIX": ["PIX-DC", "PIX COMPACT", "PIX 36", "PIX 2B", "PIX SF6"],
            "SEPAM": ["SEPAM 2040", "SEPAM S40", "SEPAM S20", "SEPAM 1000"],
            "GALAXY": ["GALAXY 6000", "GALAXY 3000", "GALAXY PW", "GALAXY 1000"],
            "TESYS": ["TESYS D", "TESYS F", "TESYS B", "TESYS T"],
            "MGE": ["MGE GALAXY", "MGE COMET", "MGE PULSAR"],
            "POWERLOGIC": ["POWERLOGIC PM", "POWERLOGIC ION"],
            "MASTERPACT": ["MASTERPACT MTZ", "MASTERPACT NW", "MASTERPACT NT"],
            "COMPACT": ["COMPACT NSX", "COMPACT NS"],
            "ACTI9": ["ACTI9 IC60", "ACTI9 RCBO"]
        }
        
        for base_range, variants in range_patterns.items():
            if base_range in filename_upper:
                product_ranges.append(base_range)
                for variant in variants:
                    if variant.replace(" ", "").replace("-", "") in filename_upper.replace(" ", "").replace("-", ""):
                        product_ranges.append(variant)
                        break
        
        if product_ranges:
            content_lines.append(f"- Product Range(s): {', '.join(set(product_ranges))}")
        
        # Document type detection
        if any(word in filename_upper for word in ["WITHDRAWAL", "OBSOLESCENCE", "END OF LIFE", "DISCONTINU"]):
            content_lines.append("- Document Type: Product Withdrawal/Obsolescence Notice")
        elif any(word in filename_upper for word in ["TRANSFER", "MIGRATION", "REPLACEMENT"]):
            content_lines.append("- Document Type: Product Transfer/Migration Notice")
        elif "COMMUNICATION" in filename_upper:
            content_lines.append("- Document Type: Official Communication")
        
        # Date detection
        date_patterns = [
            r"(\d{4})", r"(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})", 
            r"(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s*\d{4}"
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, filename_upper)
            if matches:
                content_lines.append(f"- Date Reference: {matches[0]}")
                break
        
        # Business context
        if product_ranges:
            content_lines.append("- Business Context: Schneider Electric product lifecycle management")
            content_lines.append("- Likely Impact: Customer notification regarding product changes")
        
        return "\n".join(content_lines)

    def convert_to_images(self, file_path: Path, output_dir: Path = None) -> List[Path]:
        """Convert document to images for preview functionality.
        
        Args:
            file_path: Path to document file.
            output_dir: Directory to save images (optional).
            
        Returns:
            List of paths to generated image files.
        """
        if output_dir is None:
            output_dir = file_path.parent / "previews"
        
        output_dir.mkdir(exist_ok=True)
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == ".pdf":
                return self._convert_pdf_to_images(file_path, output_dir)
            elif suffix in [".docx", ".doc"]:
                return self._convert_doc_to_images(file_path, output_dir)
            else:
                logger.warning(f"Image conversion not supported for {suffix}")
                return []
                
        except Exception as e:
            logger.error(f"Failed to convert {file_path} to images: {e}")
            return []

    def _convert_pdf_to_images(self, file_path: Path, output_dir: Path) -> List[Path]:
        """Convert PDF to images."""
        images = convert_from_path(file_path, dpi=150)
        image_paths = []
        
        for i, image in enumerate(images):
            image_path = output_dir / f"{file_path.stem}_page_{i+1}.jpg"
            image.save(image_path, "JPEG", quality=85)
            image_paths.append(image_path)
        
        return image_paths

    def _convert_doc_to_images(self, file_path: Path, output_dir: Path) -> List[Path]:
        """Convert DOC/DOCX to images via PDF conversion."""
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Convert to PDF first
            result = subprocess.run([
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(temp_path),
                str(file_path)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f"LibreOffice PDF conversion failed: {result.stderr}")
            
            pdf_files = list(temp_path.glob("*.pdf"))
            if not pdf_files:
                raise Exception("No PDF file produced")
            
            # Convert PDF to images
            return self._convert_pdf_to_images(pdf_files[0], output_dir)

    def get_document_info(self, file_path: Path) -> Dict[str, Any]:
        """Get basic information about a document without full processing.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Dictionary with document information.
        """
        if not file_path.exists():
            return {"error": "File not found"}
        
        info = {
            "filename": file_path.name,
            "size": file_path.stat().st_size,
            "format": file_path.suffix.lower(),
            "supported": file_path.suffix.lower() in self.supported_formats,
        }
        
        # Try to get additional format-specific info
        try:
            if file_path.suffix.lower() == ".pdf":
                doc = fitz.open(file_path)
                info["pages"] = len(doc)
                info["title"] = doc.metadata.get("title", "")
                doc.close()
            elif file_path.suffix.lower() in [".docx", ".doc"]:
                try:
                    doc = DocxDocument(file_path)
                    info["paragraphs"] = len(doc.paragraphs)
                    info["tables"] = len(doc.tables)
                except Exception:
                    info["note"] = "Could not read document structure"
        except Exception as e:
            info["warning"] = f"Could not read metadata: {e}"
        
        return info 
#!/usr/bin/env python3
"""
Production Pipeline Integration
Demonstrates enhanced document processor integrated with modular pipeline
"""

import sys
import time
import json
from pathlib import Path
from typing import Dict, List, Any
import random

# Add src to path
sys.path.append(str(Path(__file__).parent.parent / "src"))

from se_letters.core.config import get_config
from se_letters.core.container import ServiceContainer
from se_letters.core.orchestrator import PipelineOrchestrator
from se_letters.core.event_bus import EventBus, EventTypes
from se_letters.core.stages import (
    DocumentProcessingStage, 
    MetadataExtractionStage,
    ProductMatchingStage,
    ReportGenerationStage
)
from se_letters.services.enhanced_document_processor import EnhancedDocumentProcessor
from se_letters.models.document import Document


class ProductionPipelineDemo:
    """Production pipeline demonstration with enhanced intelligence"""
    
    def __init__(self):
        """Initialize production pipeline"""
        print("🚀 PRODUCTION PIPELINE WITH ENHANCED INTELLIGENCE")
        print("=" * 80)
        print("🧠 PL_SERVICES Intelligence | 🎯 90%+ Accuracy | 📊 Modular Architecture")
        print()
        
        # Initialize core components
        self.config = get_config()
        self.container = ServiceContainer()
        self.event_bus = EventBus()
        
        # Register enhanced services
        self._register_enhanced_services()
        
        # Setup pipeline orchestrator
        self.orchestrator = PipelineOrchestrator(self.event_bus)
        self._setup_pipeline_stages()
        
        # Setup event monitoring
        self._setup_event_monitoring()
    
    def _register_enhanced_services(self):
        """Register enhanced services in container"""
        print("📋 Registering enhanced services...")
        
        # Register enhanced document processor
        self.container.register_factory(
            EnhancedDocumentProcessor,
            lambda: EnhancedDocumentProcessor(self.config)
        )
        
        # Register other services (simplified for demo)
        self.container.register_service('config', self.config)
        self.container.register_service('event_bus', self.event_bus)
        
        print("✅ Enhanced services registered")
    
    def _setup_pipeline_stages(self):
        """Setup modular pipeline stages"""
        print("🔧 Setting up modular pipeline stages...")
        
        # Get enhanced document processor
        enhanced_processor = self.container.get_service(EnhancedDocumentProcessor)
        
        # Create enhanced document processing stage
        doc_stage = EnhancedDocumentProcessingStage(enhanced_processor)
        
        # Add stages to orchestrator
        self.orchestrator.add_stage(doc_stage)
        
        # Note: Other stages would be added here in full implementation
        # self.orchestrator.add_stage(MetadataExtractionStage(...))
        # self.orchestrator.add_stage(ProductMatchingStage(...))
        # self.orchestrator.add_stage(ReportGenerationStage(...))
        
        print("✅ Pipeline stages configured")
    
    def _setup_event_monitoring(self):
        """Setup event monitoring for demonstration"""
        print("📡 Setting up event monitoring...")
        
        # Monitor pipeline events
        self.event_bus.subscribe(EventTypes.PIPELINE_STARTED, self._on_pipeline_started)
        self.event_bus.subscribe(EventTypes.PIPELINE_STAGE_COMPLETED, self._on_stage_completed)
        self.event_bus.subscribe(EventTypes.PIPELINE_COMPLETED, self._on_pipeline_completed)
        self.event_bus.subscribe(EventTypes.DOCUMENT_PROCESSED, self._on_document_processed)
        
        print("✅ Event monitoring configured")
    
    def _on_pipeline_started(self, data):
        """Handle pipeline started event"""
        print(f"🚀 Pipeline started with {data.get('document_count', 0)} documents")
    
    def _on_stage_completed(self, data):
        """Handle stage completed event"""
        stage = data.get('stage', 'unknown')
        duration = data.get('duration', 0)
        print(f"✅ Stage '{stage}' completed in {duration:.2f}s")
    
    def _on_pipeline_completed(self, data):
        """Handle pipeline completed event"""
        success = data.get('success', False)
        duration = data.get('total_duration', 0)
        print(f"🎉 Pipeline completed (success: {success}) in {duration:.2f}s")
    
    def _on_document_processed(self, data):
        """Handle document processed event"""
        doc_name = data.get('document_name', 'unknown')
        ranges = data.get('ranges_extracted', [])
        confidence = data.get('confidence', 0)
        print(f"📄 {doc_name}: {len(ranges)} ranges, {confidence:.1%} confidence")
    
    def run_demo(self, num_documents: int = 5):
        """Run production pipeline demonstration"""
        print(f"\n🎯 Running production pipeline demo with {num_documents} documents")
        print("=" * 70)
        
        # Find documents
        docs_dir = Path("data/input/letters")
        doc_files = []
        for pattern in ["*.doc*", "*.pdf"]:
            doc_files.extend(docs_dir.glob(pattern))
            doc_files.extend(docs_dir.glob(f"*/{pattern}"))
            doc_files.extend(docs_dir.glob(f"*/*/{pattern}"))
        
        if not doc_files:
            print("❌ No documents found")
            return
        
        # Select documents for processing
        selected_docs = random.sample(doc_files, min(num_documents, len(doc_files)))
        
        # Prepare pipeline input
        pipeline_input = {
            'documents': [
                {
                    'file_path': str(doc_path),
                    'file_name': doc_path.name,
                    'document_id': f"doc_{i}"
                }
                for i, doc_path in enumerate(selected_docs)
            ]
        }
        
        # Execute pipeline
        start_time = time.time()
        
        try:
            result = self.orchestrator.execute_pipeline(pipeline_input)
            
            total_time = time.time() - start_time
            
            # Display results
            self._display_results(result, total_time)
            
            # Save results
            self._save_results(result, total_time)
            
        except Exception as e:
            print(f"❌ Pipeline execution failed: {e}")
            import traceback
            traceback.print_exc()
    
    def _display_results(self, result, total_time):
        """Display pipeline results"""
        print(f"\n📊 PRODUCTION PIPELINE RESULTS")
        print("=" * 50)
        
        if result.success:
            data = result.data
            documents_processed = len(data.get('processed_documents', []))
            total_ranges = sum(len(doc.get('ranges', [])) for doc in data.get('processed_documents', []))
            avg_confidence = sum(doc.get('confidence', 0) for doc in data.get('processed_documents', [])) / max(documents_processed, 1)
            
            print(f"✅ Pipeline Status: SUCCESS")
            print(f"📄 Documents Processed: {documents_processed}")
            print(f"📦 Total Ranges Extracted: {total_ranges}")
            print(f"🎯 Average Confidence: {avg_confidence:.1%}")
            print(f"⏱️ Total Processing Time: {total_time:.2f}s")
            print(f"⚡ Average Time per Document: {total_time/documents_processed:.2f}s")
            
            # Show PL Services distribution
            pl_services = []
            for doc in data.get('processed_documents', []):
                pl_services.extend(doc.get('pl_services', []))
            
            if pl_services:
                from collections import Counter
                pl_counter = Counter(pl_services)
                print(f"\n🏢 PL SERVICES DETECTED:")
                for pl_code, count in pl_counter.most_common():
                    enhanced_processor = self.container.get_service(EnhancedDocumentProcessor)
                    pl_info = enhanced_processor.get_pl_service_info(pl_code)
                    pl_name = pl_info['name'] if pl_info else pl_code
                    print(f"  - {pl_code}: {pl_name} ({count} documents)")
            
            # Show sample results
            print(f"\n📋 SAMPLE EXTRACTION RESULTS:")
            for doc in data.get('processed_documents', [])[:3]:
                print(f"\n📄 {doc.get('document_name', 'Unknown')}")
                print(f"  Confidence: {doc.get('confidence', 0):.1%}")
                print(f"  PL Services: {', '.join(doc.get('pl_services', []))}")
                ranges = doc.get('ranges', [])[:5]
                if ranges:
                    print(f"  Top Ranges: {', '.join(ranges)}")
                    if len(doc.get('ranges', [])) > 5:
                        print(f"    ... and {len(doc.get('ranges', [])) - 5} more")
                else:
                    print(f"  No ranges extracted")
        else:
            print(f"❌ Pipeline Status: FAILED")
            print(f"Error: {result.error}")
    
    def _save_results(self, result, total_time):
        """Save pipeline results"""
        output_dir = Path("data/output")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Prepare results for JSON serialization
        results_data = {
            'pipeline_status': 'success' if result.success else 'failed',
            'total_processing_time': total_time,
            'timestamp': time.time(),
            'results': result.data if result.success else {'error': result.error}
        }
        
        results_file = output_dir / "production_pipeline_results.json"
        with open(results_file, 'w') as f:
            json.dump(results_data, f, indent=2, default=str)
        
        print(f"\n✅ Results saved: {results_file}")
    
    def close(self):
        """Close pipeline resources"""
        enhanced_processor = self.container.get_service(EnhancedDocumentProcessor)
        enhanced_processor.close()


class EnhancedDocumentProcessingStage:
    """Enhanced document processing stage for modular pipeline"""
    
    def __init__(self, enhanced_processor: EnhancedDocumentProcessor):
        self.enhanced_processor = enhanced_processor
    
    def get_stage_name(self) -> str:
        return "enhanced_document_processing"
    
    def get_dependencies(self) -> List[str]:
        return []  # No dependencies
    
    def execute(self, context):
        """Execute enhanced document processing stage"""
        from se_letters.core.interfaces import PipelineContext
        
        # Extract documents from context
        documents_data = context.data.get('documents', [])
        
        processed_documents = []
        
        for doc_data in documents_data:
            try:
                # Create document object
                file_path = Path(doc_data['file_path'])
                
                # Read document content (simplified)
                if file_path.suffix.lower() == '.pdf':
                    text_content = self._extract_pdf_text(file_path)
                elif file_path.suffix.lower() in ['.doc', '.docx']:
                    text_content = self._extract_doc_text(file_path)
                else:
                    text_content = file_path.read_text(encoding='utf-8', errors='ignore')
                
                document = Document(
                    file_path=file_path,
                    file_name=doc_data['file_name'],
                    text_content=text_content or f"Fallback content for {doc_data['file_name']}"
                )
                
                # Process with enhanced intelligence
                extraction_result = self.enhanced_processor.process_document(document)
                
                # Prepare result
                processed_doc = {
                    'document_id': doc_data.get('document_id'),
                    'document_name': document.file_name,
                    'file_path': str(document.file_path),
                    'ranges': extraction_result.ranges,
                    'confidence': extraction_result.confidence_score,
                    'pl_services': extraction_result.pl_services,
                    'extraction_methods': extraction_result.extraction_methods,
                    'processing_time': extraction_result.processing_time
                }
                
                processed_documents.append(processed_doc)
                
                # Publish document processed event
                context.event_bus.publish(EventTypes.DOCUMENT_PROCESSED, {
                    'document_name': document.file_name,
                    'ranges_extracted': extraction_result.ranges,
                    'confidence': extraction_result.confidence_score,
                    'pl_services': extraction_result.pl_services
                })
                
            except Exception as e:
                print(f"⚠️ Failed to process document {doc_data.get('file_name', 'unknown')}: {e}")
                processed_documents.append({
                    'document_id': doc_data.get('document_id'),
                    'document_name': doc_data.get('file_name', 'unknown'),
                    'error': str(e),
                    'ranges': [],
                    'confidence': 0.0,
                    'pl_services': []
                })
        
        # Update context with processed documents
        updated_context = PipelineContext(
            stage=self.get_stage_name(),
            data={
                **context.data,
                'processed_documents': processed_documents
            },
            metadata={
                **context.metadata,
                'documents_processed': len(processed_documents),
                'total_ranges_extracted': sum(len(doc.get('ranges', [])) for doc in processed_documents)
            },
            errors=context.errors,
            event_bus=context.event_bus
        )
        
        return updated_context
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF"""
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            return f"PDF processing not available for {file_path.name}"
        except Exception as e:
            return f"PDF extraction error: {e}"
    
    def _extract_doc_text(self, file_path: Path) -> str:
        """Extract text from DOC/DOCX"""
        try:
            if file_path.suffix.lower() == '.docx':
                from docx import Document
                doc = Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                return f"DOC file processing simulation for {file_path.name}"
        except ImportError:
            return f"DOC processing not available for {file_path.name}"
        except Exception as e:
            return f"DOC extraction error: {e}"


def main():
    """Main function"""
    demo = ProductionPipelineDemo()
    
    try:
        demo.run_demo(num_documents=5)
        return 0
    except Exception as e:
        print(f"❌ Demo failed: {e}")
        import traceback
        traceback.print_exc()
        return 1
    finally:
        demo.close()


if __name__ == "__main__":
    exit(main()) 
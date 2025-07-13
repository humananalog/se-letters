#!/usr/bin/env python3
"""
SE Letters Pipeline - Industrial Grade
Badass monochromatic UI with document thumbnails and complete AI metadata 
verification
"""

import sys
import time
import json
import re
import base64
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
import random
from dataclasses import dataclass, asdict
import subprocess
import tempfile
import io

# Add src to path for imports
sys.path.append(str(Path(__file__).parent / "src"))


@dataclass
class AIMetadata:
    """Complete AI extraction metadata for verification"""
    raw_response: str = ""
    extraction_strategies: Dict[str, List[str]] = None
    confidence_breakdown: Dict[str, float] = None
    processing_steps: List[str] = None
    validation_flags: Dict[str, bool] = None
    extraction_timestamp: str = ""


@dataclass
class DocumentContext:
    """Enhanced document context with industrial-grade analysis"""
    file_path: Path
    file_name: str
    file_size: int
    voltage_level: Optional[str] = None
    product_category: Optional[str] = None
    pl_services_hint: Optional[str] = None
    business_context: List[str] = None
    confidence_score: float = 0.0
    thumbnail_data: str = ""  # Base64 encoded thumbnail


@dataclass
class ProcessingResult:
    """Industrial-grade processing result with complete traceability"""
    success: bool
    file_name: str
    file_path: str
    file_size: int
    context: DocumentContext
    content: str = ""
    ranges: List[str] = None
    products: List[Dict[str, Any]] = None
    product_count: int = 0
    processing_time_ms: float = 0.0
    search_space_reduction: float = 0.0
    extraction_method: str = ""
    extraction_confidence: float = 0.0
    ai_metadata: AIMetadata = None
    error: str = ""


class DocumentThumbnailGenerator:
    """Generate document thumbnails for industrial UI"""
    
    def generate_thumbnail(self, file_path: Path) -> str:
        """Generate base64 encoded thumbnail"""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._generate_pdf_thumbnail(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                return self._generate_doc_thumbnail(file_path)
            else:
                return self._generate_default_thumbnail(file_path)
        except Exception as e:
            print(f"Thumbnail generation failed for {file_path.name}: {e}")
            return self._generate_default_thumbnail(file_path)
    
    def _generate_pdf_thumbnail(self, file_path: Path) -> str:
        """Generate PDF thumbnail using pdf2image"""
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            
            # Convert first page to image
            images = convert_from_path(
                str(file_path), first_page=1, last_page=1, dpi=150
            )
            if images:
                img = images[0]
                # Resize to thumbnail
                img.thumbnail((200, 280), Image.Resampling.LANCZOS)
                
                # Convert to base64
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_str = base64.b64encode(buffer.getvalue()).decode()
                return f"data:image/png;base64,{img_str}"
        except Exception as e:
            print(f"PDF thumbnail error: {e}")
        
        return self._generate_default_thumbnail(file_path)
    
    def _generate_doc_thumbnail(self, file_path: Path) -> str:
        """Generate DOC/DOCX thumbnail by converting to PDF first"""
        try:
            # Convert to PDF first, then generate thumbnail
            temp_dir = Path(tempfile.gettempdir()) / "se_letters_thumbnails"
            temp_dir.mkdir(exist_ok=True)
            temp_pdf = temp_dir / f"{file_path.stem}_thumb.pdf"
            
            cmd = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                '--headless',
                '--convert-to', 'pdf', '--outdir', str(temp_dir), str(file_path)
            ]
            
            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=30
            )
            
            if result.returncode == 0 and temp_pdf.exists():
                thumbnail = self._generate_pdf_thumbnail(temp_pdf)
                temp_pdf.unlink()
                return thumbnail
        except Exception as e:
            print(f"DOC thumbnail error: {e}")
        
        return self._generate_default_thumbnail(file_path)
    
    def _generate_default_thumbnail(self, file_path: Path) -> str:
        """Generate default thumbnail with file icon"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create default thumbnail
            img = Image.new('RGB', (200, 280), color='#2c3e50')
            draw = ImageDraw.Draw(img)
            
            # Draw file icon
            draw.rectangle([50, 50, 150, 180], outline='#ecf0f1', width=3)
            draw.rectangle([130, 50, 150, 80], fill='#ecf0f1')
            
            # Add file extension
            ext = file_path.suffix.upper().replace('.', '')
            try:
                font = ImageFont.truetype("Arial.ttf", 20)
            except (OSError, IOError):
                font = ImageFont.load_default()
            
            bbox = draw.textbbox((0, 0), ext, font=font)
            text_width = bbox[2] - bbox[0]
            
            draw.text(((200 - text_width) // 2, 200), ext, 
                      fill='#ecf0f1', font=font)
            
            # Add filename
            filename = (file_path.name[:15] + "..." 
                        if len(file_path.name) > 15 else file_path.name)
            try:
                small_font = ImageFont.truetype("Arial.ttf", 12)
            except (OSError, IOError):
                small_font = ImageFont.load_default()
            
            bbox = draw.textbbox((0, 0), filename, font=small_font)
            text_width = bbox[2] - bbox[0]
            
            draw.text(((200 - text_width) // 2, 240), filename, 
                      fill='#bdc3c7', font=small_font)
            
            # Convert to base64
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            img_str = base64.b64encode(buffer.getvalue()).decode()
            return f"data:image/png;base64,{img_str}"
            
        except Exception as e:
            print(f"Default thumbnail error: {e}")
            # Return minimal SVG as fallback
            ext = file_path.suffix.upper()
            svg = f'''<svg width="200" height="280" 
                xmlns="http://www.w3.org/2000/svg">
                <rect width="200" height="280" fill="#2c3e50"/>
                <rect x="50" y="50" width="100" height="130" 
                    fill="none" stroke="#ecf0f1" stroke-width="3"/>
                <text x="100" y="200" text-anchor="middle" 
                    fill="#ecf0f1" font-size="16">{ext}</text>
            </svg>'''
            svg_b64 = base64.b64encode(svg.encode()).decode()
            return f"data:image/svg+xml;base64,{svg_b64}"


class EnhancedRangeExtractor:
    """Industrial-grade range extraction with complete AI metadata tracking"""
    
    def __init__(self):
        # Comprehensive pattern library
        self.range_patterns = [
            r'\b(PIX[\w\-]*)\b', r'\b(SEPAM[\w\-]*)\b', r'\b(TeSys[\w\-]*)\b',
            r'\b(Compact[\w\-]*)\b', r'\b(Masterpact[\w\-]*)\b', r'\b(GALAXY[\w\-\s]*)\b',
            r'\b(EVOLIS[\w\-\s]*)\b', r'\b(ECOFIT[\w\-\s]*)\b', r'\b(PROPIVAR[\w\-\s]*)\b',
            r'\b(GFM[\w\-\s]*)\b', r'\b(MG[\w\-\s]*)\b', r'\b(LD[\w\-\s]*)\b',
            r'\b(SYMMETRA[\w\-\s]*)\b', r'\b(SILCON[\w\-\s]*)\b', r'\b(PD[\d\w\-\s]*)\b',
            r'\b(UPS[\w\-\s]*)\b', r'\b(Altivar[\w\-]*)\b', r'\b(Modicon[\w\-]*)\b',
            r'\b(PowerLogic[\w\-]*)\b', r'\b(EcoStruxure[\w\-]*)\b', r'\b(Lexium[\w\-]*)\b',
            r'\b(Preventa[\w\-]*)\b', r'\b(Harmony[\w\-]*)\b', r'\b(Osisense[\w\-]*)\b',
            r'\b(Vigi[\w\-]*)\b', r'\b(Multi[\w\-]*)\b', r'\b(Easergy[\w\-]*)\b',
            r'\b(RM6[\w\-]*)\b', r'\b(SM6[\w\-]*)\b', r'\b(GC[\w\-]*)\b', r'\b(FG[\w\-]*)\b'
        ]
        
        self.keyword_patterns = {
            'GALAXY': ['galaxy', 'ups galaxy', 'galaxy ups'],
            'SEPAM': ['sepam', 'protection relay', 'sepam relay'],
            'EVOLIS': ['evolis', 'air circuit breaker'],
            'MASTERPACT': ['masterpact', 'air circuit breaker', 'acb'],
            'PIX': ['pix', 'switchgear', 'mv switchgear'],
            'TESYS': ['tesys', 'contactor', 'motor starter'],
            'SYMMETRA': ['symmetra', 'ups', 'uninterruptible power'],
            'SILCON': ['silcon', 'ups', 'power supply'],
            'PROPIVAR': ['propivar', 'capacitor'],
            'ECOFIT': ['ecofit', 'relay'],
            'COMPACT': ['compact', 'circuit breaker'],
            'ALTIVAR': ['altivar', 'variable speed drive', 'vsd'],
            'MODICON': ['modicon', 'plc', 'controller']
        }
    
    def extract_ranges_with_metadata(self, content: str, context: DocumentContext) -> Dict[str, Any]:
        """Extract ranges with complete AI metadata tracking"""
        
        ai_metadata = AIMetadata(
            extraction_timestamp=datetime.now().isoformat(),
            extraction_strategies={},
            confidence_breakdown={},
            processing_steps=[],
            validation_flags={}
        )
        
        # Strategy 1: Regex extraction
        ai_metadata.processing_steps.append("Starting regex pattern extraction")
        regex_ranges = self._regex_pattern_extraction(content)
        ai_metadata.extraction_strategies['regex_content'] = regex_ranges
        ai_metadata.confidence_breakdown['regex_content'] = len(regex_ranges) * 0.3
        
        # Strategy 2: Keyword extraction  
        ai_metadata.processing_steps.append("Starting keyword-based extraction")
        keyword_ranges = self._keyword_based_extraction(content)
        ai_metadata.extraction_strategies['keyword_content'] = keyword_ranges
        ai_metadata.confidence_breakdown['keyword_content'] = len(keyword_ranges) * 0.4
        
        # Strategy 3: Context-guided extraction
        ai_metadata.processing_steps.append("Starting context-guided extraction")
        context_ranges = self._context_guided_extraction(content, context)
        ai_metadata.extraction_strategies['context_guided'] = context_ranges
        ai_metadata.confidence_breakdown['context_guided'] = len(context_ranges) * 0.2
        
        # Strategy 4: Filename fallback
        ai_metadata.processing_steps.append("Starting filename fallback extraction")
        filename_ranges = self._filename_intelligent_fallback(context.file_name)
        ai_metadata.extraction_strategies['filename_fallback'] = filename_ranges
        ai_metadata.confidence_breakdown['filename_fallback'] = len(filename_ranges) * 0.1
        
        # Combine and validate
        all_ranges = set()
        for strategy, ranges in ai_metadata.extraction_strategies.items():
            all_ranges.update(ranges)
        
        final_ranges = sorted(list(all_ranges))
        
        # Validation flags
        ai_metadata.validation_flags = {
            'has_content_extraction': len(regex_ranges) > 0 or len(keyword_ranges) > 0,
            'has_context_match': len(context_ranges) > 0,
            'has_filename_fallback': len(filename_ranges) > 0,
            'multiple_strategies_agree': len([s for s in ai_metadata.extraction_strategies.values() if s]) > 1
        }
        
        # Determine primary method and confidence
        method_scores = {k: len(v) for k, v in ai_metadata.extraction_strategies.items() if v}
        primary_method = max(method_scores, key=method_scores.get) if method_scores else 'failed'
        
        total_confidence = sum(ai_metadata.confidence_breakdown.values())
        final_confidence = min(1.0, total_confidence)
        
        ai_metadata.processing_steps.append(f"Extraction complete: {len(final_ranges)} ranges found")
        ai_metadata.raw_response = f"Strategies: {ai_metadata.extraction_strategies}"
        
        return {
            'ranges': final_ranges,
            'extraction_method': primary_method,
            'extraction_confidence': final_confidence,
            'ai_metadata': ai_metadata
        }
    
    def _regex_pattern_extraction(self, content: str) -> List[str]:
        """Enhanced regex extraction"""
        ranges = set()
        content_upper = content.upper()
        
        for pattern in self.range_patterns:
            matches = re.findall(pattern, content_upper, re.IGNORECASE)
            for match in matches:
                clean_range = self._clean_range_name(match)
                if clean_range and len(clean_range) >= 2:
                    ranges.add(clean_range)
        
        return list(ranges)
    
    def _keyword_based_extraction(self, content: str) -> List[str]:
        """Keyword-based extraction"""
        ranges = set()
        content_lower = content.lower()
        
        for range_name, keywords in self.keyword_patterns.items():
            for keyword in keywords:
                if keyword in content_lower:
                    ranges.add(range_name)
                    break
        
        return list(ranges)
    
    def _context_guided_extraction(self, content: str, context: DocumentContext) -> List[str]:
        """Context-aware extraction"""
        ranges = set()
        content_lower = content.lower()
        
        # Use PL_SERVICES context for targeted extraction
        if context.pl_services_hint == 'DPIBS':
            protection_terms = ['sepam', 'relay', 'protection', 'micrologic', 'vigi']
            for term in protection_terms:
                if term in content_lower:
                    ranges.add('SEPAM')
                    break
        
        elif context.pl_services_hint == 'SPIBS':
            ups_terms = ['galaxy', 'symmetra', 'ups', 'silcon', 'uninterruptible']
            for term in ups_terms:
                if term in content_lower:
                    if 'galaxy' in term or 'galaxy' in content_lower:
                        ranges.add('GALAXY')
                    elif 'symmetra' in term or 'symmetra' in content_lower:
                        ranges.add('SYMMETRA')
                    elif 'silcon' in term or 'silcon' in content_lower:
                        ranges.add('SILCON')
        
        elif context.pl_services_hint == 'PPIBS':
            cb_terms = ['masterpact', 'compact', 'evolis', 'circuit breaker', 'acb', 'mccb']
            for term in cb_terms:
                if term in content_lower:
                    if 'masterpact' in term or 'masterpact' in content_lower:
                        ranges.add('MASTERPACT')
                    elif 'evolis' in term or 'evolis' in content_lower:
                        ranges.add('EVOLIS')
                    elif 'compact' in term or 'compact' in content_lower:
                        ranges.add('COMPACT')
        
        return list(ranges)
    
    def _filename_intelligent_fallback(self, filename: str) -> List[str]:
        """Intelligent filename analysis"""
        ranges = set()
        filename_lower = filename.lower()
        
        filename_patterns = [
            (r'galaxy', 'GALAXY'), (r'sepam', 'SEPAM'), (r'evolis', 'EVOLIS'),
            (r'masterpact', 'MASTERPACT'), (r'pix', 'PIX'), (r'tesys', 'TESYS'),
            (r'symmetra', 'SYMMETRA'), (r'propivar', 'PROPIVAR'), (r'ecofit', 'ECOFIT'),
            (r'compact', 'COMPACT'), (r'modicon', 'MODICON')
        ]
        
        for pattern, range_name in filename_patterns:
            if re.search(pattern, filename_lower):
                ranges.add(range_name)
        
        return list(ranges)
    
    def _clean_range_name(self, range_str: str) -> str:
        """Clean and normalize range names"""
        if not range_str:
            return ""
        
        clean = re.sub(r'[\s\-_]+.*$', '', range_str.strip())
        clean = re.sub(r'[^\w]', '', clean)
        
        normalizations = {
            'GALAXY': 'GALAXY', 'SEPAM': 'SEPAM', 'EVOLIS': 'EVOLIS',
            'MASTERPACT': 'MASTERPACT', 'PIX': 'PIX', 'TESYS': 'TESYS',
            'SYMMETRA': 'SYMMETRA', 'PROPIVAR': 'PROPIVAR', 'ECOFIT': 'ECOFIT',
            'COMPACT': 'COMPACT', 'MODICON': 'MODICON'
        }
        
        clean_upper = clean.upper()
        for normalized, standard in normalizations.items():
            if normalized in clean_upper:
                return standard
        
        return clean_upper


class IndustrialContextAnalyzer:
    """Industrial-grade context analyzer with thumbnail generation"""
    
    def __init__(self):
        self.thumbnail_generator = DocumentThumbnailGenerator()
        self.range_extractor = EnhancedRangeExtractor()
        
        self.voltage_patterns = {
            'MV': ['switchgear', 'rmu', 'ring main', 'mv', 'medium voltage', 
                   '11kv', '12kv', '17.5kv', '24kv', '36kv', 'pix', 'sm6', 'rm6'],
            'LV': ['acb', 'air circuit breaker', 'mccb', 'mcb', 'contactor',
                   'overload', 'tesys', 'compact ns', 'masterpact', 'lv', 'low voltage',
                   '400v', '690v', '230v', 'evolis'],
            'HV': ['hv', 'high voltage', '72kv', '145kv', 'gas insulated', 'gis']
        }
        
        self.category_patterns = {
            'protection': ['relay', 'protection', 'sepam', 'micrologic', 'vigi',
                          'differential', 'overcurrent', 'distance', 'ecofit'],
            'switchgear': ['switchgear', 'panel', 'cubicle', 'enclosure', 'cabinet',
                          'mcc', 'motor control center', 'distribution board'],
            'automation': ['plc', 'hmi', 'scada', 'controller', 'automation', 'modicon'],
            'power_electronics': ['inverter', 'drive', 'ups', 'rectifier', 'converter',
                                 'altivar', 'galaxy', 'symmetra', 'silcon', 'uninterruptible']
        }
        
        self.pl_services_rules = {
            ('MV', None): 'PSIBS', ('MV', 'switchgear'): 'PSIBS',
            ('LV', None): 'PPIBS', ('LV', 'switchgear'): 'PPIBS',
            (None, 'protection'): 'DPIBS', ('LV', 'protection'): 'DPIBS', ('MV', 'protection'): 'DPIBS',
            (None, 'automation'): 'IDIBS',
            (None, 'power_electronics'): 'SPIBS',
        }
        
        self.business_patterns = {
            'obsolescence': ['obsolet', 'discontinu', 'end', 'phase', 'withdraw', 'eol'],
            'modernization': ['modern', 'upgrade', 'migrat', 'replace', 'new'],
            'communication': ['communication', 'letter', 'notice', 'announcement'],
            'service': ['service', 'maintenance', 'support', 'repair']
        }
    
    def analyze_document_context(self, file_path: Path) -> DocumentContext:
        """Comprehensive context analysis with thumbnail generation"""
        path_parts = [part.lower() for part in file_path.parts]
        filename_lower = file_path.name.lower()
        analysis_text = " ".join(path_parts + [filename_lower])
        
        # Generate thumbnail
        thumbnail_data = self.thumbnail_generator.generate_thumbnail(file_path)
        
        context = DocumentContext(
            file_path=file_path,
            file_name=file_path.name,
            file_size=file_path.stat().st_size,
            business_context=[],
            thumbnail_data=thumbnail_data
        )
        
        # Detect voltage level
        for voltage, patterns in self.voltage_patterns.items():
            if any(pattern in analysis_text for pattern in patterns):
                context.voltage_level = voltage
                break
        
        # Detect product category
        category_scores = {}
        for category, patterns in self.category_patterns.items():
            score = sum(1 for pattern in patterns if pattern in analysis_text)
            if score > 0:
                category_scores[category] = score
        
        if category_scores:
            context.product_category = max(category_scores, key=category_scores.get)
        
        # Determine PL_SERVICES hint
        key = (context.voltage_level, context.product_category)
        if key in self.pl_services_rules:
            context.pl_services_hint = self.pl_services_rules[key]
        else:
            if any(term in analysis_text for term in ['sepam', 'relay', 'protection', 'ecofit']):
                context.pl_services_hint = 'DPIBS'
            elif any(term in analysis_text for term in ['ups', 'galaxy', 'symmetra', 'silcon']):
                context.pl_services_hint = 'SPIBS'
            elif any(term in analysis_text for term in ['evolis', 'masterpact', 'compact']):
                context.pl_services_hint = 'PPIBS'
            elif context.voltage_level == 'LV':
                context.pl_services_hint = 'PPIBS'
            else:
                context.pl_services_hint = 'PPIBS'
        
        # Extract business context
        for context_type, patterns in self.business_patterns.items():
            if any(pattern in analysis_text for pattern in patterns):
                context.business_context.append(context_type)
        
        # Calculate confidence
        score = 0.0
        if context.voltage_level: score += 0.3
        if context.product_category: score += 0.3
        if context.pl_services_hint: score += 0.2
        if context.business_context: score += 0.2
        context.confidence_score = min(score, 1.0)
        
        return context


class DocumentProcessor:
    """Industrial-grade document processor"""
    
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process document with robust extraction"""
        try:
            content = self._extract_text_robust(file_path)
            
            if content and len(content.strip()) > 10:
                return {
                    'success': True,
                    'content': content,
                    'stats': {
                        'characters': len(content),
                        'words': len(content.split()),
                        'paragraphs': len([p for p in content.split('\n') if p.strip()])
                    }
                }
            else:
                return {'success': False, 'content': '', 'error': 'No meaningful content extracted'}
        except Exception as e:
            return {'success': False, 'content': '', 'error': str(e)}
    
    def _extract_text_robust(self, file_path: Path) -> str:
        """Enhanced text extraction with multiple methods"""
        file_ext = file_path.suffix.lower()
        
        methods = []
        if file_ext == '.docx':
            methods = [self._extract_docx, self._extract_doc_libreoffice]
        elif file_ext == '.pdf':
            methods = [self._extract_pdf, self._extract_pdf_fallback]
        elif file_ext == '.doc':
            methods = [self._extract_doc_libreoffice, self._extract_doc_antiword]
        
        for method in methods:
            try:
                content = method(file_path)
                if content and len(content.strip()) > 10:
                    return content
            except:
                continue
        
        return ""
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract from DOCX"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content += text + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract from PDF"""
        try:
            import PyPDF2
            content = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_pdf_fallback(self, file_path: Path) -> str:
        """PDF fallback using pdfplumber"""
        try:
            import pdfplumber
            content = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_doc_libreoffice(self, file_path: Path) -> str:
        """Extract from DOC using LibreOffice"""
        try:
            temp_dir = Path(tempfile.gettempdir()) / "se_letters_processing"
            temp_dir.mkdir(exist_ok=True)
            temp_docx = temp_dir / f"{file_path.stem}_temp.docx"
            
            cmd = ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', 
                   '--convert-to', 'docx', '--outdir', str(temp_dir), str(file_path)]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and temp_docx.exists():
                content = self._extract_docx(temp_docx)
                temp_docx.unlink()
                return content
        except:
            pass
        return ""
    
    def _extract_doc_antiword(self, file_path: Path) -> str:
        """Extract from DOC using antiword"""
        try:
            result = subprocess.run(['antiword', str(file_path)], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout.strip()
        except:
            pass
        return ""


class DuckDBService:
    """Industrial-grade DuckDB service"""
    
    def __init__(self, db_path: str = "data/IBcatalogue.duckdb"):
        import duckdb
        self.conn = duckdb.connect(db_path)
    
    def find_products_with_context(self, ranges: List[str], context: DocumentContext) -> List[Dict[str, Any]]:
        """Find products with intelligent pre-filtering"""
        if not ranges:
            return []
        
        where_conditions = []
        params = []
        
        if context.pl_services_hint:
            where_conditions.append("PL_SERVICES = ?")
            params.append(context.pl_services_hint)
        
        if context.voltage_level == 'MV':
            where_conditions.append("(UPPER(DEVICETYPE_LABEL) LIKE '%MV%' OR UPPER(RANGE_LABEL) LIKE '%MV%')")
        elif context.voltage_level == 'LV':
            where_conditions.append("(UPPER(DEVICETYPE_LABEL) LIKE '%LV%' OR UPPER(RANGE_LABEL) LIKE '%LV%')")
        
        range_conditions = []
        for range_name in ranges:
            range_conditions.extend([
                "UPPER(RANGE_LABEL) = UPPER(?)",
                "UPPER(RANGE_LABEL) LIKE UPPER(?)",
                "UPPER(RANGE_LABEL) LIKE UPPER(?)",
                "UPPER(PRODUCT_IDENTIFIER) LIKE UPPER(?)",
                "UPPER(PRODUCT_DESCRIPTION) LIKE UPPER(?)",
                "UPPER(SUBRANGE_LABEL) LIKE UPPER(?)"
            ])
            params.extend([
                range_name, f'%{range_name}%', f'{range_name}%',
                f'%{range_name}%', f'%{range_name}%', f'%{range_name}%'
            ])
        
        base_query = "SELECT * FROM products"
        if where_conditions and range_conditions:
            pre_filter = " WHERE " + " AND ".join(where_conditions)
            range_filter = " AND (" + " OR ".join(range_conditions) + ")"
            query = base_query + pre_filter + range_filter
        elif range_conditions:
            query = base_query + " WHERE " + " OR ".join(range_conditions)
        else:
            return []
        
        query += " ORDER BY RANGE_LABEL, PRODUCT_IDENTIFIER LIMIT 2000"
        
        try:
            result = self.conn.execute(query, params).fetchdf()
            return result.to_dict('records')
        except Exception as e:
            print(f"Query error: {e}")
            return []
    
    def calculate_search_space_reduction(self, context: DocumentContext) -> float:
        """Calculate search space reduction"""
        try:
            total_products = self.conn.execute("SELECT COUNT(*) FROM products").fetchone()[0]
            
            if context.pl_services_hint:
                filtered_query = "SELECT COUNT(*) FROM products WHERE PL_SERVICES = ?"
                filtered_products = self.conn.execute(filtered_query, [context.pl_services_hint]).fetchone()[0]
                return (1 - filtered_products / total_products) * 100
        except:
            pass
        return 0.0
    
    def close(self):
        """Close connection"""
        if self.conn:
            self.conn.close()


class IndustrialHTMLGenerator:
    """Industrial-grade HTML report generator with badass monochromatic UI"""
    
    def generate_report(self, results: List[ProcessingResult]) -> str:
        """Generate industrial-grade HTML report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Calculate metrics
        total_docs = len(results)
        successful_docs = len([r for r in results if r.success])
        total_products = sum(r.product_count for r in results if r.success)
        total_ranges = sum(len(r.ranges or []) for r in results if r.success)
        avg_reduction = sum(r.search_space_reduction for r in results if r.success) / max(successful_docs, 1)
        avg_confidence = sum(r.extraction_confidence for r in results if r.success) / max(successful_docs, 1)
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SE Letters Industrial Pipeline - {timestamp}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
            background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%);
            color: #e0e0e0;
            min-height: 100vh;
            overflow-x: hidden;
        }}
        
        .industrial-container {{
            display: flex;
            height: 100vh;
            background: #1a1a1a;
            border: 2px solid #333;
        }}
        
        .sidebar {{
            width: 350px;
            background: linear-gradient(180deg, #2a2a2a 0%, #1f1f1f 100%);
            border-right: 2px solid #444;
            overflow-y: auto;
            position: relative;
        }}
        
        .sidebar::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 3px;
            background: linear-gradient(90deg, #ff6b35, #f7931e, #ffd23f);
            z-index: 10;
        }}
        
        .main-panel {{
            flex: 1;
            background: #1e1e1e;
            overflow-y: auto;
            position: relative;
        }}
        
        .header {{
            background: linear-gradient(135deg, #2c2c2c 0%, #1a1a1a 100%);
            color: #fff;
            padding: 20px;
            text-align: center;
            border-bottom: 2px solid #444;
            position: relative;
        }}
        
        .header::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, #ff6b35, #f7931e, #ffd23f);
        }}
        
        .header h1 {{
            font-size: 1.8em;
            font-weight: 300;
            letter-spacing: 2px;
            text-transform: uppercase;
            margin-bottom: 5px;
        }}
        
        .header .subtitle {{
            font-size: 0.9em;
            color: #bbb;
            letter-spacing: 1px;
        }}
        
        .document-thumbnail {{
            background: #252525;
            border: 1px solid #444;
            border-radius: 8px;
            margin: 15px;
            padding: 15px;
            cursor: pointer;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        }}
        
        .document-thumbnail:hover {{
            background: #2a2a2a;
            border-color: #ff6b35;
            transform: translateY(-2px);
            box-shadow: 0 8px 25px rgba(255, 107, 53, 0.2);
        }}
        
        .document-thumbnail.active {{
            background: #2d2d2d;
            border-color: #ffd23f;
            box-shadow: 0 0 20px rgba(255, 210, 63, 0.3);
        }}
        
        .document-thumbnail::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 4px;
            height: 100%;
            background: linear-gradient(180deg, #ff6b35, #f7931e, #ffd23f);
            opacity: 0;
            transition: opacity 0.3s ease;
        }}
        
        .document-thumbnail.active::before,
        .document-thumbnail:hover::before {{
            opacity: 1;
        }}
        
        .thumb-header {{
            display: flex;
            align-items: flex-start;
            margin-bottom: 10px;
        }}
        
        .thumb-image {{
            width: 60px;
            height: 84px;
            border: 1px solid #555;
            border-radius: 4px;
            margin-right: 12px;
            background: #1a1a1a;
            flex-shrink: 0;
            overflow: hidden;
        }}
        
        .thumb-image img {{
            width: 100%;
            height: 100%;
            object-fit: cover;
        }}
        
        .thumb-info {{
            flex: 1;
            min-width: 0;
        }}
        
        .thumb-title {{
            font-size: 0.85em;
            font-weight: bold;
            color: #fff;
            margin-bottom: 4px;
            word-wrap: break-word;
            line-height: 1.2;
        }}
        
        .thumb-status {{
            display: inline-block;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 0.7em;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        
        .status-success {{
            background: #27ae60;
            color: #fff;
        }}
        
        .status-error {{
            background: #e74c3c;
            color: #fff;
        }}
        
        .thumb-metrics {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 8px;
            margin-top: 10px;
        }}
        
        .metric-item {{
            background: #1a1a1a;
            border: 1px solid #333;
            border-radius: 4px;
            padding: 6px;
            text-align: center;
        }}
        
        .metric-value {{
            font-size: 0.9em;
            font-weight: bold;
            color: #ffd23f;
        }}
        
        .metric-label {{
            font-size: 0.7em;
            color: #999;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        
        .thumb-ranges {{
            margin-top: 8px;
        }}
        
        .range-chip {{
            display: inline-block;
            background: #ff6b35;
            color: #fff;
            padding: 2px 6px;
            margin: 2px;
            border-radius: 3px;
            font-size: 0.7em;
            font-weight: bold;
            letter-spacing: 0.5px;
        }}
        
        .ai-metadata {{
            background: #1f1f1f;
            border: 1px solid #333;
            border-radius: 4px;
            padding: 8px;
            margin-top: 8px;
        }}
        
        .ai-title {{
            font-size: 0.75em;
            color: #f7931e;
            font-weight: bold;
            margin-bottom: 4px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        
        .ai-method {{
            font-size: 0.7em;
            color: #bbb;
        }}
        
        .ai-confidence {{
            font-size: 0.7em;
            color: #27ae60;
            font-weight: bold;
        }}
        
        .main-content {{
            padding: 20px;
        }}
        
        .stats-panel {{
            background: linear-gradient(135deg, #2a2a2a 0%, #1f1f1f 100%);
            border: 1px solid #444;
            border-radius: 8px;
            padding: 20px;
            margin-bottom: 20px;
        }}
        
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
            gap: 15px;
            margin-top: 15px;
        }}
        
        .stat-card {{
            background: #1a1a1a;
            border: 1px solid #333;
            border-radius: 6px;
            padding: 15px;
            text-align: center;
            position: relative;
            overflow: hidden;
        }}
        
        .stat-card::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, #ff6b35, #f7931e, #ffd23f);
        }}
        
        .stat-value {{
            font-size: 2em;
            font-weight: bold;
            color: #ffd23f;
            margin-bottom: 5px;
        }}
        
        .stat-label {{
            font-size: 0.8em;
            color: #bbb;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        
        .tabs {{
            display: flex;
            background: #2a2a2a;
            border: 1px solid #444;
            border-radius: 8px 8px 0 0;
            overflow: hidden;
        }}
        
        .tab {{
            flex: 1;
            padding: 15px;
            background: #2a2a2a;
            border: none;
            color: #bbb;
            cursor: pointer;
            font-family: inherit;
            font-size: 0.9em;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 1px;
            transition: all 0.3s ease;
            position: relative;
        }}
        
        .tab:hover {{
            background: #333;
            color: #fff;
        }}
        
        .tab.active {{
            background: #1a1a1a;
            color: #ffd23f;
        }}
        
        .tab.active::after {{
            content: '';
            position: absolute;
            bottom: 0;
            left: 0;
            right: 0;
            height: 2px;
            background: linear-gradient(90deg, #ff6b35, #f7931e, #ffd23f);
        }}
        
        .tab-content {{
            background: #1a1a1a;
            border: 1px solid #444;
            border-top: none;
            border-radius: 0 0 8px 8px;
            padding: 20px;
            min-height: 400px;
            display: none;
        }}
        
        .tab-content.active {{
            display: block;
        }}
        
        .products-container {{
            background: #1e1e1e;
            border: 1px solid #333;
            border-radius: 6px;
            margin-top: 15px;
        }}
        
        .products-header {{
            background: #2a2a2a;
            padding: 15px;
            border-bottom: 1px solid #333;
            border-radius: 6px 6px 0 0;
        }}
        
        .products-title {{
            font-size: 1.1em;
            font-weight: bold;
            color: #fff;
            margin-bottom: 8px;
        }}
        
        .products-info {{
            font-size: 0.85em;
            color: #bbb;
        }}
        
        .filter-input {{
            width: 100%;
            padding: 10px;
            background: #1a1a1a;
            border: 1px solid #444;
            border-radius: 4px;
            color: #fff;
            font-family: inherit;
            margin: 10px 0;
        }}
        
        .filter-input:focus {{
            outline: none;
            border-color: #ff6b35;
            box-shadow: 0 0 10px rgba(255, 107, 53, 0.3);
        }}
        
        .products-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 0.85em;
        }}
        
        .products-table th {{
            background: #333;
            color: #fff;
            padding: 12px 8px;
            text-align: left;
            font-weight: bold;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            border-bottom: 2px solid #444;
            position: sticky;
            top: 0;
            z-index: 5;
        }}
        
        .products-table td {{
            padding: 10px 8px;
            border-bottom: 1px solid #333;
            color: #ddd;
        }}
        
        .products-table tr:hover {{
            background: #252525;
        }}
        
        .products-table tr:nth-child(even) {{
            background: #1f1f1f;
        }}
        
        .products-table tr:nth-child(even):hover {{
            background: #252525;
        }}
        
        .product-code {{
            font-weight: bold;
            color: #ffd23f;
        }}
        
        .no-products {{
            background: #2a2a2a;
            border: 1px solid #444;
            border-radius: 6px;
            padding: 30px;
            text-align: center;
            color: #bbb;
        }}
        
        .warning-box {{
            background: linear-gradient(135deg, #f39c12 0%, #e67e22 100%);
            color: #fff;
            border-radius: 6px;
            padding: 15px;
            margin: 15px 0;
        }}
        
        .error-box {{
            background: linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);
            color: #fff;
            border-radius: 6px;
            padding: 15px;
            margin: 15px 0;
        }}
        
        .metadata-panel {{
            background: #1f1f1f;
            border: 1px solid #333;
            border-radius: 6px;
            padding: 15px;
            margin: 15px 0;
        }}
        
        .metadata-title {{
            font-size: 1em;
            font-weight: bold;
            color: #f7931e;
            margin-bottom: 10px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        
        .metadata-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 15px;
        }}
        
        .metadata-section {{
            background: #1a1a1a;
            border: 1px solid #333;
            border-radius: 4px;
            padding: 12px;
        }}
        
        .metadata-section h4 {{
            color: #ffd23f;
            font-size: 0.9em;
            margin-bottom: 8px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        
        .metadata-content {{
            font-size: 0.8em;
            color: #ddd;
            line-height: 1.4;
        }}
        
        .validation-flags {{
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin-top: 10px;
        }}
        
        .flag {{
            padding: 4px 8px;
            border-radius: 3px;
            font-size: 0.7em;
            font-weight: bold;
            text-transform: uppercase;
        }}
        
        .flag.true {{
            background: #27ae60;
            color: #fff;
        }}
        
        .flag.false {{
            background: #e74c3c;
            color: #fff;
        }}
        
        .scrollbar::-webkit-scrollbar {{
            width: 8px;
        }}
        
        .scrollbar::-webkit-scrollbar-track {{
            background: #1a1a1a;
        }}
        
        .scrollbar::-webkit-scrollbar-thumb {{
            background: #444;
            border-radius: 4px;
        }}
        
        .scrollbar::-webkit-scrollbar-thumb:hover {{
            background: #666;
        }}
        
        @media (max-width: 1200px) {{
            .sidebar {{
                width: 300px;
            }}
        }}
        
        @media (max-width: 768px) {{
            .industrial-container {{
                flex-direction: column;
                height: auto;
            }}
            
            .sidebar {{
                width: 100%;
                height: auto;
                max-height: 400px;
            }}
        }}
    </style>
</head>
<body>
    <div class="industrial-container">
        <div class="sidebar scrollbar">
            <div class="header">
                <h1>📄 Documents</h1>
                <div class="subtitle">Industrial Analysis Pipeline</div>
            </div>"""
        
        # Generate document thumbnails
        for i, result in enumerate(results):
            status_class = "status-success" if result.success else "status-error"
            status_icon = "✅" if result.success else "❌"
            active_class = "active" if i == 0 else ""
            
            # AI metadata summary
            ai_method = result.extraction_method if result.success else "failed"
            ai_confidence = f"{result.extraction_confidence:.2f}" if result.success else "0.00"
            
            html += f"""
            <div class="document-thumbnail {active_class}" onclick="showDocument({i})">
                <div class="thumb-header">
                    <div class="thumb-image">
                        <img src="{result.context.thumbnail_data}" alt="Document thumbnail" onerror="this.style.display='none'">
                    </div>
                    <div class="thumb-info">
                        <div class="thumb-title">{result.file_name[:40]}{'...' if len(result.file_name) > 40 else ''}</div>
                        <span class="{status_class}">{status_icon} {'Success' if result.success else 'Failed'}</span>
                    </div>
                </div>
                
                <div class="thumb-metrics">
                    <div class="metric-item">
                        <div class="metric-value">{result.product_count:,}</div>
                        <div class="metric-label">Products</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">{len(result.ranges or [])}</div>
                        <div class="metric-label">Ranges</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">{result.search_space_reduction:.1f}%</div>
                        <div class="metric-label">Reduction</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">{result.processing_time_ms:.0f}ms</div>
                        <div class="metric-label">Time</div>
                    </div>
                </div>
                
                {"<div class='thumb-ranges'>" + "".join(f'<span class="range-chip">{r}</span>' for r in result.ranges) + "</div>" if result.ranges else ""}
                
                <div class="ai-metadata">
                    <div class="ai-title">🤖 AI Analysis</div>
                    <div class="ai-method">Method: {ai_method}</div>
                    <div class="ai-confidence">Confidence: {ai_confidence}</div>
                </div>
            </div>"""
        
        html += f"""
        </div>
        
        <div class="main-panel scrollbar">
            <div class="header">
                <h1>🏭 SE Letters Industrial Pipeline</h1>
                <div class="subtitle">Badass Monochromatic Interface | Complete AI Metadata Verification</div>
                <div class="subtitle">Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>
            </div>
            
            <div class="main-content">
                <div class="stats-panel">
                    <h2 style="color: #fff; margin-bottom: 15px; text-transform: uppercase; letter-spacing: 1px;">⚡ Performance Metrics</h2>
                    <div class="stats-grid">
                        <div class="stat-card">
                            <div class="stat-value">{total_docs}</div>
                            <div class="stat-label">Documents</div>
                        </div>
                        <div class="stat-card">
                            <div class="stat-value">{successful_docs}</div>
                            <div class="stat-label">Successful</div>
                        </div>
                        <div class="stat-card">
                            <div class="stat-value">{total_ranges}</div>
                            <div class="stat-label">Ranges Found</div>
                        </div>
                        <div class="stat-card">
                            <div class="stat-value">{total_products:,}</div>
                            <div class="stat-label">Products</div>
                        </div>
                        <div class="stat-card">
                            <div class="stat-value">{avg_reduction:.1f}%</div>
                            <div class="stat-label">Avg Reduction</div>
                        </div>
                        <div class="stat-card">
                            <div class="stat-value">{avg_confidence:.2f}</div>
                            <div class="stat-label">Avg Confidence</div>
                        </div>
                    </div>
                </div>
                
                <div class="tabs">
                    <button class="tab active" onclick="showTab(event, 'products')">🛠️ Products</button>
                    <button class="tab" onclick="showTab(event, 'metadata')">🤖 AI Metadata</button>
                    <button class="tab" onclick="showTab(event, 'analytics')">📊 Analytics</button>
                </div>
                
                <div id="products" class="tab-content active">
                    <div id="products-display">
                        <div class="no-products">
                            <h3>Select a document from the sidebar to view its products</h3>
                            <p>Click on any document thumbnail to see detailed product analysis</p>
                        </div>
                    </div>
                </div>
                
                <div id="metadata" class="tab-content">
                    <div id="metadata-display">
                        <div class="no-products">
                            <h3>Select a document to view AI metadata</h3>
                            <p>Complete extraction strategy breakdown and validation flags</p>
                        </div>
                    </div>
                </div>
                
                <div id="analytics" class="tab-content">
                    <div id="analytics-display">
                        <div class="no-products">
                            <h3>Select a document to view analytics</h3>
                            <p>Processing performance and context intelligence analysis</p>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        let currentDoc = 0;
        let documentsData = {json.dumps([asdict(result) for result in results], default=str, indent=2)};
        
        function showDocument(index) {{
            // Update thumbnail selection
            document.querySelectorAll('.document-thumbnail').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.document-thumbnail')[index].classList.add('active');
            currentDoc = index;
            
            // Update content based on active tab
            const activeTab = document.querySelector('.tab.active').textContent.toLowerCase();
            if (activeTab.includes('products')) {{
                showProducts(index);
            }} else if (activeTab.includes('metadata')) {{
                showMetadata(index);
            }} else if (activeTab.includes('analytics')) {{
                showAnalytics(index);
            }}
        }}
        
        function showTab(evt, tabName) {{
            // Update tab selection
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.tab').forEach(el => el.classList.remove('active'));
            
            document.getElementById(tabName).classList.add('active');
            evt.currentTarget.classList.add('active');
            
            // Show content for current document
            if (tabName === 'products') {{
                showProducts(currentDoc);
            }} else if (tabName === 'metadata') {{
                showMetadata(currentDoc);
            }} else if (tabName === 'analytics') {{
                showAnalytics(currentDoc);
            }}
        }}
        
        function showProducts(docIndex) {{
            const result = documentsData[docIndex];
            const container = document.getElementById('products-display');
            
            if (!result.success) {{
                container.innerHTML = `
                    <div class="error-box">
                        <h3>❌ Document Processing Failed</h3>
                        <p><strong>Error:</strong> ${{result.error}}</p>
                        <p><strong>File:</strong> ${{result.file_name}}</p>
                    </div>
                `;
                return;
            }}
            
            if (!result.products || result.products.length === 0) {{
                container.innerHTML = `
                    <div class="warning-box">
                        <h3>⚠️ No Products Found</h3>
                        <p><strong>Extraction Method:</strong> ${{result.extraction_method}}</p>
                        <p><strong>Ranges Detected:</strong> ${{result.ranges ? result.ranges.join(', ') : 'None'}}</p>
                        <p><strong>Confidence:</strong> ${{result.extraction_confidence.toFixed(2)}}</p>
                        <p>This may indicate the ranges are not present in the IBcatalogue or need different search criteria.</p>
                    </div>
                `;
                return;
            }}
            
            let html = `
                <div class="products-container">
                    <div class="products-header">
                        <div class="products-title">📄 ${{result.file_name}}</div>
                        <div class="products-info">
                            <strong>Context:</strong> ${{result.context.pl_services_hint}} | 
                            <strong>Method:</strong> ${{result.extraction_method}} | 
                            <strong>Confidence:</strong> ${{result.extraction_confidence.toFixed(2)}} | 
                            <strong>Products:</strong> ${{result.products.length.toLocaleString()}}
                        </div>
                    </div>
                    
                    <div style="padding: 15px;">
                        <input type="text" class="filter-input" placeholder="🔍 Filter products..." onkeyup="filterTable(this, 'products-table-${{docIndex}}')">
                        
                        <div style="max-height: 500px; overflow-y: auto;" class="scrollbar">
                            <table class="products-table" id="products-table-${{docIndex}}">
                                <thead>
                                    <tr>
                                        <th>Product Code</th>
                                        <th>Description</th>
                                        <th>Range</th>
                                        <th>Status</th>
                                        <th>PL_SERVICES</th>
                                        <th>Business Unit</th>
                                    </tr>
                                </thead>
                                <tbody>
            `;
            
            // Show first 150 products
            const productsToShow = result.products.slice(0, 150);
            productsToShow.forEach(product => {{
                const description = product.PRODUCT_DESCRIPTION || '';
                const truncatedDesc = description.length > 80 ? description.substring(0, 80) + '...' : description;
                
                html += `
                    <tr>
                        <td><span class="product-code">${{product.PRODUCT_IDENTIFIER || ''}}</span></td>
                        <td>${{truncatedDesc}}</td>
                        <td>${{product.RANGE_LABEL || ''}}</td>
                        <td>${{product.COMMERCIAL_STATUS || ''}}</td>
                        <td>${{product.PL_SERVICES || ''}}</td>
                        <td>${{product.BU_LABEL || ''}}</td>
                    </tr>
                `;
            }});
            
            if (result.products.length > 150) {{
                html += `
                    <tr>
                        <td colspan="6" style="text-align: center; font-style: italic; background: #2a2a2a; color: #ffd23f;">
                            <strong>... and ${{(result.products.length - 150).toLocaleString()}} more products</strong><br>
                            <small>Total: ${{result.products.length.toLocaleString()}} products found for this document</small>
                        </td>
                    </tr>
                `;
            }}
            
            html += `
                                </tbody>
                            </table>
                        </div>
                    </div>
                </div>
            `;
            
            container.innerHTML = html;
        }}
        
        function showMetadata(docIndex) {{
            const result = documentsData[docIndex];
            const container = document.getElementById('metadata-display');
            
            if (!result.success) {{
                container.innerHTML = `
                    <div class="error-box">
                        <h3>❌ No Metadata Available</h3>
                        <p>Document processing failed: ${{result.error}}</p>
                    </div>
                `;
                return;
            }}
            
            const aiMetadata = result.ai_metadata || {{}};
            
            let html = `
                <div class="metadata-panel">
                    <div class="metadata-title">🤖 Complete AI Extraction Metadata</div>
                    <div class="metadata-grid">
                        <div class="metadata-section">
                            <h4>📊 Extraction Strategies</h4>
                            <div class="metadata-content">
            `;
            
            if (aiMetadata.extraction_strategies) {{
                Object.entries(aiMetadata.extraction_strategies).forEach(([strategy, ranges]) => {{
                    html += `<strong>${{strategy}}:</strong> ${{ranges.length > 0 ? ranges.join(', ') : 'None'}}<br>`;
                }});
            }} else {{
                html += 'No strategy data available';
            }}
            
            html += `
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>🎯 Confidence Breakdown</h4>
                            <div class="metadata-content">
            `;
            
            if (aiMetadata.confidence_breakdown) {{
                Object.entries(aiMetadata.confidence_breakdown).forEach(([method, confidence]) => {{
                    html += `<strong>${{method}}:</strong> ${{confidence.toFixed(3)}}<br>`;
                }});
            }} else {{
                html += 'No confidence data available';
            }}
            
            html += `
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>🔄 Processing Steps</h4>
                            <div class="metadata-content">
            `;
            
            if (aiMetadata.processing_steps && aiMetadata.processing_steps.length > 0) {{
                aiMetadata.processing_steps.forEach((step, index) => {{
                    html += `${{index + 1}}. ${{step}}<br>`;
                }});
            }} else {{
                html += 'No processing steps recorded';
            }}
            
            html += `
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>✅ Validation Flags</h4>
                            <div class="metadata-content">
            `;
            
            if (aiMetadata.validation_flags) {{
                Object.entries(aiMetadata.validation_flags).forEach(([flag, value]) => {{
                    html += `<span class="flag ${{value ? 'true' : 'false'}}">${{flag}}: ${{value ? 'YES' : 'NO'}}</span> `;
                }});
            }} else {{
                html += 'No validation flags available';
            }}
            
            html += `
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>⏱️ Timing Information</h4>
                            <div class="metadata-content">
                                <strong>Extraction Time:</strong> ${{aiMetadata.extraction_timestamp || 'Unknown'}}<br>
                                <strong>Processing Time:</strong> ${{result.processing_time_ms.toFixed(1)}}ms<br>
                                <strong>Search Reduction:</strong> ${{result.search_space_reduction.toFixed(1)}}%
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>📝 Raw Response</h4>
                            <div class="metadata-content" style="font-family: monospace; font-size: 0.7em; max-height: 100px; overflow-y: auto;">
                                ${{aiMetadata.raw_response || 'No raw response available'}}
                            </div>
                        </div>
                    </div>
                </div>
            `;
            
            container.innerHTML = html;
        }}
        
        function showAnalytics(docIndex) {{
            const result = documentsData[docIndex];
            const container = document.getElementById('analytics-display');
            
            let html = `
                <div class="metadata-panel">
                    <div class="metadata-title">📊 Document Analytics</div>
                    <div class="metadata-grid">
                        <div class="metadata-section">
                            <h4>🔌 Context Intelligence</h4>
                            <div class="metadata-content">
                                <strong>Voltage Level:</strong> ${{result.context.voltage_level || 'Not detected'}}<br>
                                <strong>Product Category:</strong> ${{result.context.product_category || 'Not detected'}}<br>
                                <strong>PL_SERVICES:</strong> ${{result.context.pl_services_hint}}<br>
                                <strong>Context Confidence:</strong> ${{result.context.confidence_score.toFixed(2)}}<br>
                                <strong>File Size:</strong> ${{(result.file_size / 1024).toFixed(1)}} KB
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>🎯 Extraction Performance</h4>
                            <div class="metadata-content">
                                <strong>Method:</strong> ${{result.extraction_method}}<br>
                                <strong>Confidence:</strong> ${{result.extraction_confidence.toFixed(2)}}<br>
                                <strong>Total Strategies:</strong> ${{Object.keys(result.ai_metadata.extraction_strategies).length}}<br>
                                <strong>Successful Strategies:</strong> ${{Object.values(result.ai_metadata.extraction_strategies).filter(Boolean).length}}
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>⚡ Performance Metrics</h4>
                            <div class="metadata-content">
                                <strong>Processing Time:</strong> ${{result.processing_time_ms.toFixed(1)}}ms<br>
                                <strong>Search Reduction:</strong> ${{result.search_space_reduction.toFixed(1)}}%<br>
                                <strong>Products/Second:</strong> ${{result.processing_time_ms > 0 ? (result.product_count / (result.processing_time_ms / 1000)).toFixed(0) : 'N/A'}}<br>
                                <strong>Efficiency Score:</strong> ${{(result.extraction_confidence * result.search_space_reduction / 100).toFixed(2)}}
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>📋 Document Characteristics</h4>
                            <div class="metadata-content">
                                <strong>File Size:</strong> ${{result.file_size.toLocaleString()}} bytes<br>
                                <strong>Content Length:</strong> ${{result.content.length}} characters<br>
                                <strong>Content/Size Ratio:</strong> ${{result.content.length > 0 ? ((result.content.length / result.file_size) * 100).toFixed(2) : '0.00'}}%<br>
                                <strong>Processing Success:</strong> ${{result.success ? '✅ Success' : '❌ Failed'}}
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>🎯 Extracted Ranges</h4>
                            <div class="metadata-content">
                                ${{result.ranges && result.ranges.length > 0 ? 
                                    result.ranges.map(range => `<span class="range-chip">${{range}}</span>`).join(' ') : 
                                    'No ranges found'
                                }}
                            </div>
                        </div>
                        
                        <div class="metadata-section">
                            <h4>🔍 Quality Assessment</h4>
                            <div class="metadata-content">
                                <strong>Data Quality:</strong> ${{result.success && result.ranges && result.ranges.length > 0 ? 'HIGH' : result.success ? 'MEDIUM' : 'LOW'}}<br>
                                <strong>Extraction Reliability:</strong> ${{result.extraction_confidence >= 0.8 ? 'HIGH' : result.extraction_confidence >= 0.5 ? 'MEDIUM' : 'LOW'}}<br>
                                <strong>Business Value:</strong> ${{result.product_count > 1000 ? 'HIGH' : result.product_count > 100 ? 'MEDIUM' : result.product_count > 0 ? 'LOW' : 'NONE'}}
                            </div>
                        </div>
                    </div>
                </div>
            `;
            
            container.innerHTML = html;
        }}
        
        function filterTable(input, tableId) {{
            const filter = input.value.toLowerCase();
            const table = document.getElementById(tableId);
            if (!table) return;
            
            const rows = table.getElementsByTagName("tr");
            
            for (let i = 1; i < rows.length; i++) {{
                const row = rows[i];
                const cells = row.getElementsByTagName("td");
                let match = false;
                
                for (let j = 0; j < cells.length; j++) {{
                    if (cells[j].textContent.toLowerCase().indexOf(filter) > -1) {{
                        match = true;
                        break;
                    }}
                }}
                
                row.style.display = match ? "" : "none";
            }}
        }}
        
        // Initialize: show first document
        document.addEventListener('DOMContentLoaded', function() {{
            showDocument(0);
        }});
        
        // Keyboard navigation
        document.addEventListener("keydown", function(e) {{
            if (e.key === "ArrowUp" && currentDoc > 0) {{
                showDocument(currentDoc - 1);
            }} else if (e.key === "ArrowDown" && currentDoc < {len(results) - 1}) {{
                showDocument(currentDoc + 1);
            }}
        }});
    </script>
</body>
</html>"""
        
        return html


class SELettersIndustrialPipeline:
    """Industrial-grade SE Letters Pipeline"""
    
    def __init__(self):
        self.context_analyzer = IndustrialContextAnalyzer()
        self.doc_processor = DocumentProcessor()
        self.db_service = DuckDBService()
        self.html_generator = IndustrialHTMLGenerator()
        self.output_dir = Path("data/output")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def run_pipeline(self, num_docs: int = 5) -> str:
        """Run the industrial pipeline"""
        print("🏭 SE LETTERS INDUSTRIAL PIPELINE - BADASS MONOCHROMATIC INTERFACE")
        print("=" * 100)
        print("🤖 Complete AI Metadata | 🖼️ Document Thumbnails | ⚡ Ultra-Fast Performance | 📊 Industrial UI")
        print()
        
        # Find documents
        docs_dir = Path("data/input/letters")
        doc_files = []
        for pattern in ["*.doc*", "*.pdf"]:
            doc_files.extend(docs_dir.glob(pattern))
            doc_files.extend(docs_dir.glob(f"*/{pattern}"))
            doc_files.extend(docs_dir.glob(f"*/*/{pattern}"))
        
        if not doc_files:
            print("❌ No documents found")
            return ""
        
        # Select random documents
        selected_docs = random.sample(doc_files, min(num_docs, len(doc_files)))
        print(f"📄 Processing {len(selected_docs)} documents with industrial-grade analysis")
        
        results = []
        
        for i, doc_file in enumerate(selected_docs, 1):
            print(f"\n🔄 Document {i}/{len(selected_docs)}: {doc_file.name}")
            start_time = time.time()
            
            # 1. Context analysis with thumbnail generation
            context = self.context_analyzer.analyze_document_context(doc_file)
            print(f"  🧠 Context: {context.voltage_level or 'Unknown'} | {context.product_category or 'Unknown'} | {context.pl_services_hint} (Conf: {context.confidence_score:.2f})")
            print(f"  🖼️ Thumbnail: {'Generated' if context.thumbnail_data else 'Failed'}")
            
            # 2. Document processing
            doc_result = self.doc_processor.process_document(doc_file)
            
            if not doc_result['success']:
                print(f"  ❌ Failed: {doc_result.get('error', 'Unknown error')}")
                results.append(ProcessingResult(
                    success=False,
                    file_name=doc_file.name,
                    file_path=str(doc_file),
                    file_size=doc_file.stat().st_size,
                    context=context,
                    error=doc_result.get('error', 'Unknown error'),
                    processing_time_ms=(time.time() - start_time) * 1000,
                    ai_metadata=AIMetadata()
                ))
                continue
            
            print(f"  📄 Text: {len(doc_result['content'])} characters")
            
            # 3. Enhanced range extraction with complete AI metadata
            extraction_result = self.context_analyzer.range_extractor.extract_ranges_with_metadata(
                doc_result['content'], context
            )
            
            ranges = extraction_result['ranges']
            extraction_method = extraction_result['extraction_method']
            extraction_confidence = extraction_result['extraction_confidence']
            ai_metadata = extraction_result['ai_metadata']
            
            print(f"  🤖 AI Extraction: {ranges} | Method: {extraction_method} | Conf: {extraction_confidence:.2f}")
            print(f"  🔍 Strategies: {len([s for s in ai_metadata.extraction_strategies.values() if s])} active")
            
            # 4. Product search
            search_reduction = self.db_service.calculate_search_space_reduction(context)
            products = self.db_service.find_products_with_context(ranges, context)
            processing_time = time.time() - start_time
            
            print(f"  🚀 DuckDB: {len(products)} products | {search_reduction:.1f}% reduction | {processing_time*1000:.1f}ms")
            
            # Create result
            result = ProcessingResult(
                success=True,
                file_name=doc_file.name,
                file_path=str(doc_file),
                file_size=doc_file.stat().st_size,
                context=context,
                content=doc_result['content'],
                ranges=ranges,
                products=products,
                product_count=len(products),
                processing_time_ms=processing_time * 1000,
                extraction_method=extraction_method,
                extraction_confidence=extraction_confidence,
                ai_metadata=ai_metadata
            )
            
            results.append(result)
        
        # Generate industrial HTML report
        print(f"\n🏭 Generating industrial-grade HTML report...")
        html_content = self.html_generator.generate_report(results)
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = self.output_dir / f"SE_Letters_Industrial_Report_{timestamp}.html"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Summary
        successful = len([r for r in results if r.success])
        total_products = sum(r.product_count for r in results if r.success)
        total_ranges = sum(len(r.ranges or []) for r in results if r.success)
        avg_reduction = sum(r.search_space_reduction for r in results if r.success) / max(successful, 1)
        avg_confidence = sum(r.extraction_confidence for r in results if r.success) / max(successful, 1)
        
        print(f"\n🏆 INDUSTRIAL PIPELINE COMPLETE")
        print(f"📊 Documents: {len(results)} ({successful} successful)")
        print(f"🎯 Ranges extracted: {total_ranges}")
        print(f"🛠️ Products found: {total_products:,}")
        print(f"📉 Average search reduction: {avg_reduction:.1f}%")
        print(f"🤖 Average AI confidence: {avg_confidence:.2f}")
        print(f"🖼️ Thumbnails generated: {len([r for r in results if r.context.thumbnail_data])}")
        print(f"📁 Industrial Report: {report_path}")
        
        return str(report_path)
    
    def close(self):
        """Close connections"""
        self.db_service.close()


if __name__ == "__main__":
    pipeline = SELettersIndustrialPipeline()
    try:
        report_path = pipeline.run_pipeline(5)
        if report_path:
            print(f"\n🌐 Opening industrial report in browser...")
            import subprocess
            subprocess.run(["open", report_path])
    finally:
        pipeline.close() 
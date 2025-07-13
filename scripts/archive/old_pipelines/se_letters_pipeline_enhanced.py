#!/usr/bin/env python3
"""
SE Letters Pipeline - Enhanced Version with Expanded Range Detection
Addresses critical 80% range extraction failure rate with comprehensive patterns
"""

import sys
import time
import json
import yaml
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional, Set
import random
from dataclasses import dataclass
import subprocess
import tempfile

# Add src to path for imports
sys.path.append(str(Path(__file__).parent / "src"))


@dataclass
class DocumentContext:
    """Document context with intelligent pre-filtering hints"""
    file_path: Path
    file_name: str
    voltage_level: Optional[str] = None
    product_category: Optional[str] = None
    pl_services_hint: Optional[str] = None
    business_context: List[str] = None
    confidence_score: float = 0.0


@dataclass
class ProcessingResult:
    """Complete processing result"""
    success: bool
    file_name: str
    file_path: str
    file_size: int
    context: DocumentContext
    content: str = ""
    ranges: List[str] = None
    products: List[Dict[str, Any]] = None
    product_count: int = 0
    processing_time_ms: float = 0.0
    search_space_reduction: float = 0.0
    extraction_method: str = ""
    extraction_confidence: float = 0.0
    error: str = ""


class EnhancedRangeExtractor:
    """Enhanced range extraction with multiple strategies"""
    
    def __init__(self):
        # CRITICAL: Expanded pattern library to address 80% failure rate
        self.range_patterns = [
            # Original patterns
            r'\b(PIX[\w\-]*)\b',
            r'\b(SEPAM[\w\-]*)\b',
            r'\b(TeSys[\w\-]*)\b',
            r'\b(Compact[\w\-]*)\b',
            r'\b(Masterpact[\w\-]*)\b',
            r'\b(Altivar[\w\-]*)\b',
            r'\b(Modicon[\w\-]*)\b',
            r'\b(PowerLogic[\w\-]*)\b',
            r'\b(EcoStruxure[\w\-]*)\b',
            r'\b(Lexium[\w\-]*)\b',
            r'\b(Preventa[\w\-]*)\b',
            r'\b(Harmony[\w\-]*)\b',
            r'\b(Osisense[\w\-]*)\b',
            r'\b(Vigi[\w\-]*)\b',
            r'\b(Multi[\w\-]*)\b',
            r'\b(Easergy[\w\-]*)\b',
            r'\b(RM6[\w\-]*)\b',
            r'\b(SM6[\w\-]*)\b',
            r'\b(GC[\w\-]*)\b',
            r'\b(FG[\w\-]*)\b',
            
            # MISSING PATTERNS - CRITICAL ADDITIONS
            r'\b(GALAXY[\w\-\s]*)\b',
            r'\b(EVOLIS[\w\-\s]*)\b',
            r'\b(ECOFIT[\w\-\s]*)\b',
            r'\b(PROPIVAR[\w\-\s]*)\b',
            r'\b(GFM[\w\-\s]*)\b',
            r'\b(MG[\w\-\s]*)\b',
            r'\b(LD[\w\-\s]*)\b',
            r'\b(SYMMETRA[\w\-\s]*)\b',
            r'\b(SILCON[\w\-\s]*)\b',
            r'\b(PD[\d\w\-\s]*)\b',
            r'\b(UPS[\w\-\s]*)\b',
            
            # Product code patterns
            r'\b(SEPAM\s+[A-Z]+[\d\w]*)\b',
            r'\b(Galaxy\s+[\d\w]+)\b',
            r'\b(Symmetra\s+[\d\w]+)\b',
            r'\b(Masterpact\s+[\w\d]+)\b',
            
            # Protection relay patterns
            r'\b(PROTECTION\s+RELAY[\w\-\s]*)\b',
            r'\b(MV\s+PROTECTION[\w\-\s]*)\b',
            r'\b(RELAY[\w\-\s]*)\b'
        ]
        
        # Keyword-based patterns for fallback
        self.keyword_patterns = {
            'GALAXY': ['galaxy', 'ups galaxy', 'galaxy ups'],
            'SEPAM': ['sepam', 'protection relay', 'sepam relay'],
            'EVOLIS': ['evolis', 'air circuit breaker'],
            'MASTERPACT': ['masterpact', 'air circuit breaker', 'acb'],
            'PIX': ['pix', 'switchgear', 'mv switchgear'],
            'TESYS': ['tesys', 'contactor', 'motor starter'],
            'SYMMETRA': ['symmetra', 'ups', 'uninterruptible power'],
            'SILCON': ['silcon', 'ups', 'power supply'],
            'PROPIVAR': ['propivar', 'capacitor'],
            'ECOFIT': ['ecofit', 'relay'],
            'COMPACT': ['compact', 'circuit breaker'],
            'ALTIVAR': ['altivar', 'variable speed drive', 'vsd'],
            'MODICON': ['modicon', 'plc', 'controller']
        }
        
        # Filename intelligence patterns
        self.filename_patterns = [
            r'(?i)\b(galaxy[\s\d\w]*)\b',
            r'(?i)\b(sepam[\s\d\w]*)\b',
            r'(?i)\b(evolis[\s\d\w]*)\b',
            r'(?i)\b(masterpact[\s\d\w]*)\b',
            r'(?i)\b(pix[\s\d\w]*)\b',
            r'(?i)\b(tesys[\s\d\w]*)\b',
            r'(?i)\b(symmetra[\s\d\w]*)\b',
            r'(?i)\b(propivar[\s\d\w]*)\b',
            r'(?i)\b(ecofit[\s\d\w]*)\b',
            r'(?i)\b(gfm[\s\d\w]*)\b',
            r'(?i)\b(mg[\s\d\w]*)\b',
            r'(?i)\b(ld[\s\d\w]*)\b'
        ]
    
    def extract_ranges_comprehensive(self, content: str, context: DocumentContext) -> Dict[str, Any]:
        """Multi-strategy comprehensive range extraction"""
        
        strategies = {
            'regex_content': self._regex_pattern_extraction(content),
            'keyword_content': self._keyword_based_extraction(content),
            'context_guided': self._context_guided_extraction(content, context),
            'filename_fallback': self._filename_intelligent_fallback(context.file_name)
        }
        
        # Combine and validate ranges
        all_ranges = set()
        method_scores = {}
        
        for method, ranges in strategies.items():
            if ranges:
                all_ranges.update(ranges)
                method_scores[method] = len(ranges)
        
        # Determine extraction method and confidence
        if all_ranges:
            primary_method = max(method_scores, key=method_scores.get) if method_scores else 'none'
            confidence = min(1.0, len(all_ranges) * 0.3 + (0.2 if 'content' in primary_method else 0.1))
        else:
            primary_method = 'failed'
            confidence = 0.0
        
        return {
            'ranges': sorted(list(all_ranges)),
            'extraction_method': primary_method,
            'extraction_confidence': confidence,
            'strategy_results': strategies
        }
    
    def _regex_pattern_extraction(self, content: str) -> List[str]:
        """Enhanced regex pattern extraction"""
        ranges = set()
        content_upper = content.upper()
        
        for pattern in self.range_patterns:
            matches = re.findall(pattern, content_upper, re.IGNORECASE)
            for match in matches:
                clean_range = self._clean_range_name(match)
                if clean_range and len(clean_range) >= 2:
                    ranges.add(clean_range)
        
        return list(ranges)
    
    def _keyword_based_extraction(self, content: str) -> List[str]:
        """Keyword-based extraction for missed patterns"""
        ranges = set()
        content_lower = content.lower()
        
        for range_name, keywords in self.keyword_patterns.items():
            for keyword in keywords:
                if keyword in content_lower:
                    ranges.add(range_name)
                    break
        
        return list(ranges)
    
    def _context_guided_extraction(self, content: str, context: DocumentContext) -> List[str]:
        """Context-aware extraction using intelligent hints"""
        ranges = set()
        content_lower = content.lower()
        
        # Use PL_SERVICES context for targeted extraction
        if context.pl_services_hint == 'DPIBS':
            # Digital Power - focus on protection
            protection_terms = ['sepam', 'relay', 'protection', 'micrologic', 'vigi']
            for term in protection_terms:
                if term in content_lower:
                    ranges.add('SEPAM')
                    break
        
        elif context.pl_services_hint == 'SPIBS':
            # Secure Power - focus on UPS
            ups_terms = ['galaxy', 'symmetra', 'ups', 'silcon', 'uninterruptible']
            for term in ups_terms:
                if term in content_lower:
                    if 'galaxy' in term or 'galaxy' in content_lower:
                        ranges.add('GALAXY')
                    elif 'symmetra' in term or 'symmetra' in content_lower:
                        ranges.add('SYMMETRA')
                    elif 'silcon' in term or 'silcon' in content_lower:
                        ranges.add('SILCON')
        
        elif context.pl_services_hint == 'PPIBS':
            # Power Products - focus on circuit breakers
            cb_terms = ['masterpact', 'compact', 'evolis', 'circuit breaker', 'acb', 'mccb']
            for term in cb_terms:
                if term in content_lower:
                    if 'masterpact' in term or 'masterpact' in content_lower:
                        ranges.add('MASTERPACT')
                    elif 'evolis' in term or 'evolis' in content_lower:
                        ranges.add('EVOLIS')
                    elif 'compact' in term or 'compact' in content_lower:
                        ranges.add('COMPACT')
        
        # Voltage level context
        if context.voltage_level == 'MV':
            mv_terms = ['pix', 'sm6', 'rm6', 'switchgear']
            for term in mv_terms:
                if term in content_lower:
                    ranges.add('PIX')
                    break
        
        return list(ranges)
    
    def _filename_intelligent_fallback(self, filename: str) -> List[str]:
        """Intelligent filename analysis as fallback"""
        ranges = set()
        filename_lower = filename.lower()
        
        for pattern in self.filename_patterns:
            matches = re.findall(pattern, filename_lower)
            for match in matches:
                clean_range = self._clean_range_name(match)
                if clean_range and len(clean_range) >= 2:
                    ranges.add(clean_range.upper())
        
        return list(ranges)
    
    def _clean_range_name(self, range_str: str) -> str:
        """Clean and normalize range names"""
        if not range_str:
            return ""
        
        # Remove common suffixes and clean
        clean = re.sub(r'[\s\-_]+.*$', '', range_str.strip())
        clean = re.sub(r'[^\w]', '', clean)
        
        # Normalize known variations
        normalizations = {
            'GALAXY': 'GALAXY',
            'SEPAM': 'SEPAM',
            'EVOLIS': 'EVOLIS',
            'MASTERPACT': 'MASTERPACT',
            'PIX': 'PIX',
            'TESYS': 'TESYS',
            'SYMMETRA': 'SYMMETRA',
            'PROPIVAR': 'PROPIVAR',
            'ECOFIT': 'ECOFIT',
            'COMPACT': 'COMPACT'
        }
        
        clean_upper = clean.upper()
        for normalized, standard in normalizations.items():
            if normalized in clean_upper:
                return standard
        
        return clean_upper


class IntelligentContextAnalyzer:
    """Analyzes document context for smart pre-filtering"""
    
    def __init__(self):
        self.voltage_patterns = {
            'MV': ['switchgear', 'rmu', 'ring main', 'mv', 'medium voltage', 
                   '11kv', '12kv', '17.5kv', '24kv', '36kv', 'pix', 'sm6', 'rm6'],
            'LV': ['acb', 'air circuit breaker', 'mccb', 'mcb', 'contactor',
                   'overload', 'tesys', 'compact ns', 'masterpact', 'lv', 'low voltage',
                   '400v', '690v', '230v', 'evolis'],
            'HV': ['hv', 'high voltage', '72kv', '145kv', 'gas insulated', 'gis']
        }
        
        self.category_patterns = {
            'protection': ['relay', 'protection', 'sepam', 'micrologic', 'vigi',
                          'differential', 'overcurrent', 'distance', 'ecofit'],
            'switchgear': ['switchgear', 'panel', 'cubicle', 'enclosure', 'cabinet',
                          'mcc', 'motor control center', 'distribution board'],
            'automation': ['plc', 'hmi', 'scada', 'controller', 'automation', 'modicon'],
            'power_electronics': ['inverter', 'drive', 'ups', 'rectifier', 'converter',
                                 'altivar', 'galaxy', 'symmetra', 'silcon', 'uninterruptible']
        }
        
        self.pl_services_rules = {
            ('MV', None): 'PSIBS', ('MV', 'switchgear'): 'PSIBS',
            ('LV', None): 'PPIBS', ('LV', 'switchgear'): 'PPIBS',
            (None, 'protection'): 'DPIBS', ('LV', 'protection'): 'DPIBS', ('MV', 'protection'): 'DPIBS',
            (None, 'automation'): 'IDIBS',
            (None, 'power_electronics'): 'SPIBS',
        }
        
        self.business_patterns = {
            'obsolescence': ['obsolet', 'discontinu', 'end', 'phase', 'withdraw', 'eol'],
            'modernization': ['modern', 'upgrade', 'migrat', 'replace', 'new'],
            'communication': ['communication', 'letter', 'notice', 'announcement'],
            'service': ['service', 'maintenance', 'support', 'repair']
        }
        
        self.range_extractor = EnhancedRangeExtractor()
    
    def analyze_document_context(self, file_path: Path) -> DocumentContext:
        """Analyze document for intelligent context"""
        path_parts = [part.lower() for part in file_path.parts]
        filename_lower = file_path.name.lower()
        analysis_text = " ".join(path_parts + [filename_lower])
        
        context = DocumentContext(
            file_path=file_path,
            file_name=file_path.name,
            business_context=[]
        )
        
        # Detect voltage level
        for voltage, patterns in self.voltage_patterns.items():
            if any(pattern in analysis_text for pattern in patterns):
                context.voltage_level = voltage
                break
        
        # Detect product category
        category_scores = {}
        for category, patterns in self.category_patterns.items():
            score = sum(1 for pattern in patterns if pattern in analysis_text)
            if score > 0:
                category_scores[category] = score
        
        if category_scores:
            context.product_category = max(category_scores, key=category_scores.get)
        
        # Determine PL_SERVICES hint
        key = (context.voltage_level, context.product_category)
        if key in self.pl_services_rules:
            context.pl_services_hint = self.pl_services_rules[key]
        else:
            # Check partial rules and specific patterns
            if any(term in analysis_text for term in ['sepam', 'relay', 'protection', 'ecofit']):
                context.pl_services_hint = 'DPIBS'
            elif any(term in analysis_text for term in ['ups', 'galaxy', 'symmetra', 'silcon']):
                context.pl_services_hint = 'SPIBS'
            elif any(term in analysis_text for term in ['evolis', 'masterpact', 'compact']):
                context.pl_services_hint = 'PPIBS'
            elif context.voltage_level == 'LV':
                context.pl_services_hint = 'PPIBS'
            else:
                context.pl_services_hint = 'PPIBS'  # Default
        
        # Extract business context
        for context_type, patterns in self.business_patterns.items():
            if any(pattern in analysis_text for pattern in patterns):
                context.business_context.append(context_type)
        
        # Calculate confidence
        score = 0.0
        if context.voltage_level: score += 0.3
        if context.product_category: score += 0.3
        if context.pl_services_hint: score += 0.2
        if context.business_context: score += 0.2
        context.confidence_score = min(score, 1.0)
        
        return context


class DocumentProcessor:
    """Enhanced document processor with robust text extraction"""
    
    def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process document with enhanced extraction methods"""
        try:
            content = self._extract_text_robust(file_path)
            
            if content and len(content.strip()) > 10:
                return {
                    'success': True,
                    'content': content,
                    'stats': {
                        'characters': len(content),
                        'words': len(content.split()),
                        'paragraphs': len([p for p in content.split('\n') if p.strip()])
                    }
                }
            else:
                return {'success': False, 'content': '', 'error': 'No meaningful content extracted'}
        except Exception as e:
            return {'success': False, 'content': '', 'error': str(e)}
    
    def _extract_text_robust(self, file_path: Path) -> str:
        """Enhanced text extraction with multiple fallback methods"""
        file_ext = file_path.suffix.lower()
        
        methods = []
        if file_ext == '.docx':
            methods = [self._extract_docx, self._extract_doc_libreoffice]
        elif file_ext == '.pdf':
            methods = [self._extract_pdf, self._extract_pdf_fallback]
        elif file_ext == '.doc':
            methods = [self._extract_doc_libreoffice, self._extract_doc_antiword, self._extract_doc_textract]
        
        # Try each method in sequence
        for method in methods:
            try:
                content = method(file_path)
                if content and len(content.strip()) > 10:
                    return content
            except:
                continue
        
        return ""
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract from DOCX"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content += text + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract from PDF"""
        try:
            import PyPDF2
            content = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_pdf_fallback(self, file_path: Path) -> str:
        """PDF fallback using pdfplumber"""
        try:
            import pdfplumber
            content = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            return content.strip()
        except:
            return ""
    
    def _extract_doc_libreoffice(self, file_path: Path) -> str:
        """Extract from DOC using LibreOffice"""
        try:
            temp_dir = Path(tempfile.gettempdir()) / "se_letters_processing"
            temp_dir.mkdir(exist_ok=True)
            temp_docx = temp_dir / f"{file_path.stem}_temp.docx"
            
            cmd = ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--headless', 
                   '--convert-to', 'docx', '--outdir', str(temp_dir), str(file_path)]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0 and temp_docx.exists():
                content = self._extract_docx(temp_docx)
                temp_docx.unlink()
                return content
        except:
            pass
        return ""
    
    def _extract_doc_antiword(self, file_path: Path) -> str:
        """Extract from DOC using antiword"""
        try:
            result = subprocess.run(['antiword', str(file_path)], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout.strip()
        except:
            pass
        return ""
    
    def _extract_doc_textract(self, file_path: Path) -> str:
        """Extract from DOC using textract"""
        try:
            import textract
            content = textract.process(str(file_path)).decode('utf-8')
            return content.strip()
        except:
            pass
        return ""


class DuckDBService:
    """Ultra-fast DuckDB service with intelligent pre-filtering"""
    
    def __init__(self, db_path: str = "data/IBcatalogue.duckdb"):
        import duckdb
        self.conn = duckdb.connect(db_path)
    
    def find_products_with_context(self, ranges: List[str], context: DocumentContext) -> List[Dict[str, Any]]:
        """Find products with intelligent pre-filtering"""
        if not ranges:
            return []
        
        # Build pre-filter conditions
        where_conditions = []
        params = []
        
        if context.pl_services_hint:
            where_conditions.append("PL_SERVICES = ?")
            params.append(context.pl_services_hint)
        
        if context.voltage_level == 'MV':
            where_conditions.append("(UPPER(DEVICETYPE_LABEL) LIKE '%MV%' OR UPPER(RANGE_LABEL) LIKE '%MV%')")
        elif context.voltage_level == 'LV':
            where_conditions.append("(UPPER(DEVICETYPE_LABEL) LIKE '%LV%' OR UPPER(RANGE_LABEL) LIKE '%LV%')")
        
        # Enhanced range matching with fuzzy logic
        range_conditions = []
        for range_name in ranges:
            range_conditions.extend([
                "UPPER(RANGE_LABEL) = UPPER(?)",
                "UPPER(RANGE_LABEL) LIKE UPPER(?)",
                "UPPER(RANGE_LABEL) LIKE UPPER(?)",
                "UPPER(PRODUCT_IDENTIFIER) LIKE UPPER(?)",
                "UPPER(PRODUCT_DESCRIPTION) LIKE UPPER(?)",
                "UPPER(SUBRANGE_LABEL) LIKE UPPER(?)"
            ])
            params.extend([
                range_name, 
                f'%{range_name}%', 
                f'{range_name}%',
                f'%{range_name}%',
                f'%{range_name}%',
                f'%{range_name}%'
            ])
        
        # Combine conditions
        base_query = "SELECT * FROM products"
        if where_conditions and range_conditions:
            pre_filter = " WHERE " + " AND ".join(where_conditions)
            range_filter = " AND (" + " OR ".join(range_conditions) + ")"
            query = base_query + pre_filter + range_filter
        elif range_conditions:
            query = base_query + " WHERE " + " OR ".join(range_conditions)
        else:
            return []
        
        query += " ORDER BY RANGE_LABEL, PRODUCT_IDENTIFIER LIMIT 2000"
        
        try:
            result = self.conn.execute(query, params).fetchdf()
            return result.to_dict('records')
        except Exception as e:
            print(f"Query error: {e}")
            return []
    
    def calculate_search_space_reduction(self, context: DocumentContext) -> float:
        """Calculate search space reduction percentage"""
        try:
            total_products = self.conn.execute("SELECT COUNT(*) FROM products").fetchone()[0]
            
            if context.pl_services_hint:
                filtered_query = "SELECT COUNT(*) FROM products WHERE PL_SERVICES = ?"
                filtered_products = self.conn.execute(filtered_query, [context.pl_services_hint]).fetchone()[0]
                return (1 - filtered_products / total_products) * 100
        except:
            pass
        return 0.0
    
    def close(self):
        """Close connection"""
        if self.conn:
            self.conn.close()


class HTMLReportGenerator:
    """Generates comprehensive HTML reports with enhanced metrics"""
    
    def generate_report(self, results: List[ProcessingResult]) -> str:
        """Generate enhanced HTML report with extraction analytics"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Calculate summary stats
        total_docs = len(results)
        successful_docs = len([r for r in results if r.success])
        total_products = sum(r.product_count for r in results if r.success)
        total_ranges = sum(len(r.ranges or []) for r in results if r.success)
        avg_reduction = sum(r.search_space_reduction for r in results if r.success) / max(successful_docs, 1)
        avg_confidence = sum(r.extraction_confidence for r in results if r.success) / max(successful_docs, 1)
        
        # Extraction method breakdown
        method_counts = {}
        for r in results:
            if r.success and r.extraction_method:
                method_counts[r.extraction_method] = method_counts.get(r.extraction_method, 0) + 1
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>SE Letters Enhanced Pipeline Report - {timestamp}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 0;
            padding: 0;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
        }}
        .main-container {{
            display: flex;
            height: 100vh;
            max-width: 1600px;
            margin: 0 auto;
            box-shadow: 0 0 50px rgba(0,0,0,0.3);
        }}
        .left-panel {{
            width: 400px;
            background: white;
            overflow-y: auto;
            padding: 20px;
            border-right: 3px solid #3498db;
        }}
        .right-panel {{
            flex: 1;
            overflow-y: auto;
            padding: 20px;
            background: white;
        }}
        .header {{
            background: linear-gradient(135deg, #2c3e50 0%, #3498db 100%);
            color: white;
            padding: 30px;
            text-align: center;
            margin: -20px -20px 20px -20px;
        }}
        .header h1 {{
            margin: 0;
            font-size: 2.2em;
            font-weight: 300;
        }}
        .document-card {{
            background: #f8f9fa;
            border-radius: 10px;
            padding: 15px;
            margin-bottom: 15px;
            cursor: pointer;
            transition: all 0.3s ease;
            border-left: 4px solid #3498db;
        }}
        .document-card:hover {{
            background: #e3f2fd;
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        }}
        .document-card.active {{
            background: #e8f5e8;
            border-left-color: #27ae60;
        }}
        .doc-header {{
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 8px;
        }}
        .context-info {{
            background: #e3f2fd;
            padding: 8px;
            border-radius: 5px;
            margin: 8px 0;
            font-size: 0.9em;
        }}
        .extraction-info {{
            background: #f0f8f0;
            padding: 8px;
            border-radius: 5px;
            margin: 8px 0;
            font-size: 0.85em;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(2, 1fr);
            gap: 8px;
            margin-top: 8px;
        }}
        .stat-item {{
            text-align: center;
            padding: 8px;
            background: white;
            border-radius: 5px;
            font-size: 0.85em;
        }}
        .summary-section {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        .metrics-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(120px, 1fr));
            gap: 15px;
            margin-top: 15px;
        }}
        .metric-card {{
            background: white;
            padding: 15px;
            border-radius: 8px;
            text-align: center;
            border-left: 4px solid #3498db;
        }}
        .metric-card h3 {{
            margin: 0;
            color: #3498db;
            font-size: 1.8em;
        }}
        .tabs {{
            display: flex;
            background: #e9ecef;
            border-radius: 8px;
            overflow: hidden;
            margin-bottom: 20px;
        }}
        .tab {{
            flex: 1;
            padding: 12px;
            background: #e9ecef;
            border: none;
            cursor: pointer;
            font-weight: 500;
            transition: all 0.3s ease;
        }}
        .tab:hover {{
            background: #dee2e6;
        }}
        .tab.active {{
            background: #3498db;
            color: white;
        }}
        .tab-content {{
            display: none;
        }}
        .tab-content.active {{
            display: block;
        }}
        .products-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
            font-size: 0.9em;
        }}
        .products-table th,
        .products-table td {{
            padding: 10px;
            border: 1px solid #ddd;
            text-align: left;
        }}
        .products-table th {{
            background: #3498db;
            color: white;
            position: sticky;
            top: 0;
        }}
        .products-table tr:hover {{
            background: #f8f9fa;
        }}
        .status-success {{ background: #d4edda; color: #155724; padding: 4px 8px; border-radius: 3px; }}
        .status-error {{ background: #f8d7da; color: #721c24; padding: 4px 8px; border-radius: 3px; }}
        .range-tag {{ background: #007bff; color: white; padding: 2px 6px; margin: 2px; border-radius: 3px; font-size: 0.8em; }}
        .method-badge {{ background: #17a2b8; color: white; padding: 2px 6px; border-radius: 3px; font-size: 0.8em; }}
        .confidence-badge {{ background: #ffc107; color: #212529; padding: 2px 6px; border-radius: 3px; font-size: 0.8em; }}
        .filter-input {{ width: 100%; padding: 8px; margin: 10px 0; border: 1px solid #ddd; border-radius: 4px; }}
        .reduction-badge {{ background: #28a745; color: white; padding: 2px 6px; border-radius: 10px; font-size: 0.8em; }}
        .performance-highlight {{ background: #fff3cd; padding: 15px; border-radius: 8px; border-left: 4px solid #ffc107; margin: 15px 0; }}
        .enhancement-notice {{ background: #d1ecf1; padding: 15px; border-radius: 8px; border-left: 4px solid #17a2b8; margin: 15px 0; }}
    </style>
</head>
<body>
    <div class="main-container">
        <div class="left-panel">
            <div class="header">
                <h1>📄 Documents</h1>
                <p>Enhanced Multi-Strategy Extraction</p>
            </div>"""
        
        # Add document cards with enhanced information
        for i, result in enumerate(results):
            status_icon = "✅" if result.success else "❌"
            active_class = "active" if i == 0 else ""
            
            html += f"""
            <div class="document-card {active_class}" onclick="showDocument({i})">
                <div class="doc-header">{status_icon} {result.file_name[:35]}{'...' if len(result.file_name) > 35 else ''}</div>
                
                <div class="context-info">
                    <strong>🧠 Intelligence:</strong><br>
                    Voltage: {result.context.voltage_level or 'Unknown'}<br>
                    Category: {result.context.product_category or 'Unknown'}<br>
                    PL_SERVICES: {result.context.pl_services_hint}<br>
                    Confidence: {result.context.confidence_score:.2f}
                </div>
                
                <div class="extraction-info">
                    <strong>🎯 Extraction:</strong><br>
                    Method: {result.extraction_method}<br>
                    Confidence: {result.extraction_confidence:.2f}<br>
                    Ranges: {len(result.ranges or [])}
                </div>
                
                <div class="stats-grid">
                    <div class="stat-item">
                        <strong>{result.product_count:,}</strong><br>Products
                    </div>
                    <div class="stat-item">
                        <strong>{result.search_space_reduction:.1f}%</strong><br>Reduction
                    </div>
                    <div class="stat-item">
                        <strong>{len(result.ranges or [])}</strong><br>Ranges
                    </div>
                    <div class="stat-item">
                        <strong>{result.processing_time_ms:.0f}ms</strong><br>Time
                    </div>
                </div>
                
                {"<div style='margin-top: 8px;'>" + "".join(f'<span class="range-tag">{r}</span>' for r in result.ranges) + "</div>" if result.ranges else ""}
            </div>"""
        
        html += f"""
        </div>
        
        <div class="right-panel">
            <div class="header">
                <h1>🚀 SE Letters Enhanced Pipeline</h1>
                <p>Multi-Strategy • Context-Aware • Ultra-Fast</p>
                <p>Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
            </div>
            
            <div class="enhancement-notice">
                <strong>🎯 Enhanced Multi-Strategy Extraction:</strong><br>
                Comprehensive pattern library (30+ ranges) • Context-guided detection • Filename fallback intelligence<br>
                <strong>Success Rate:</strong> {(successful_docs/total_docs)*100:.1f}% | <strong>Avg Confidence:</strong> {avg_confidence:.2f} | <strong>Total Ranges:</strong> {total_ranges}
            </div>
            
            <div class="summary-section">
                <h2>📊 Executive Summary</h2>
                <div class="metrics-grid">
                    <div class="metric-card">
                        <h3>{total_docs}</h3>
                        <p>Documents</p>
                    </div>
                    <div class="metric-card">
                        <h3>{successful_docs}</h3>
                        <p>Successful</p>
                    </div>
                    <div class="metric-card">
                        <h3>{total_ranges}</h3>
                        <p>Ranges Found</p>
                    </div>
                    <div class="metric-card">
                        <h3>{total_products:,}</h3>
                        <p>Products</p>
                    </div>
                    <div class="metric-card">
                        <h3>{avg_reduction:.1f}%</h3>
                        <p>Avg Reduction</p>
                    </div>
                    <div class="metric-card">
                        <h3>{avg_confidence:.2f}</h3>
                        <p>Avg Confidence</p>
                    </div>
                </div>
                
                <div class="performance-highlight">
                    <strong>🎯 Enhanced Extraction Results:</strong>
                    <ul>
                        <li>✅ Expanded pattern library covering 30+ Schneider Electric ranges</li>
                        <li>🚀 Multi-strategy extraction: regex + keywords + context + filename</li>
                        <li>🧠 Context-aware intelligence with PL_SERVICES targeting</li>
                        <li>📊 Extraction confidence scoring and method attribution</li>
                        <li>🏗️ Robust fallback mechanisms for maximum coverage</li>
                    </ul>
                    <strong>Method Breakdown:</strong> {' | '.join([f'{method}: {count}' for method, count in method_counts.items()])}
                </div>
            </div>
            
            <div class="tabs">
                <button class="tab active" onclick="showTab(event, 'overview')">📊 Overview</button>
                <button class="tab" onclick="showTab(event, 'products')">🛠️ Products</button>
                <button class="tab" onclick="showTab(event, 'analytics')">📈 Analytics</button>
            </div>
            
            <div id="overview" class="tab-content active">
                <h3>📋 Processing Overview</h3>
                <table class="products-table">
                    <thead>
                        <tr>
                            <th>Document</th>
                            <th>Status</th>
                            <th>Intelligence</th>
                            <th>Extraction</th>
                            <th>Ranges</th>
                            <th>Products</th>
                            <th>Time</th>
                        </tr>
                    </thead>
                    <tbody>"""
        
        for result in results:
            status_class = "status-success" if result.success else "status-error"
            status_text = "✅ Success" if result.success else "❌ Failed"
            intelligence = f"{result.context.pl_services_hint} (Conf: {result.context.confidence_score:.2f})"
            extraction = f"<span class='method-badge'>{result.extraction_method}</span> <span class='confidence-badge'>{result.extraction_confidence:.2f}</span>"
            ranges_text = ", ".join(result.ranges or []) if result.ranges else "None"
            
            html += f"""
                        <tr>
                            <td>{result.file_name[:25]}{'...' if len(result.file_name) > 25 else ''}</td>
                            <td><span class="{status_class}">{status_text}</span></td>
                            <td>{intelligence}</td>
                            <td>{extraction}</td>
                            <td>{ranges_text}</td>
                            <td>{result.product_count:,}</td>
                            <td>{result.processing_time_ms:.0f}ms</td>
                        </tr>"""
        
        html += """
                    </tbody>
                </table>
            </div>
            
            <div id="products" class="tab-content">
                <h3>🛠️ Product Matches (Enhanced DuckDB)</h3>
                <div id="products-container">
                    <p>Select a document from the left panel to view its product matches.</p>
                </div>"""
        
        # Generate hidden product sections for each document
        for i, result in enumerate(results):
            if result.success and result.products:
                html += f"""
                <div id="products-doc-{i}" class="document-products" style="display: none;">
                    <h4>📄 {result.file_name}</h4>
                    <p><strong>Context:</strong> {result.context.pl_services_hint} | <strong>Extraction:</strong> {result.extraction_method} ({result.extraction_confidence:.2f}) | <strong>Products:</strong> {len(result.products):,}</p>
                    
                    <input type="text" class="filter-input" placeholder="🔍 Filter products..." onkeyup="filterTable(this, 'table{i}')">
                    
                    <div style="max-height: 400px; overflow-y: auto;">
                        <table class="products-table" id="table{i}">
                            <thead>
                                <tr>
                                    <th>Product Code</th>
                                    <th>Description</th>
                                    <th>Range</th>
                                    <th>Status</th>
                                    <th>PL_SERVICES</th>
                                    <th>Business Unit</th>
                                </tr>
                            </thead>
                            <tbody>"""
                
                for product in result.products[:100]:  # Show more products for better analysis
                    html += f"""
                                <tr>
                                    <td><strong>{product.get('PRODUCT_IDENTIFIER', '')}</strong></td>
                                    <td>{product.get('PRODUCT_DESCRIPTION', '')[:80]}{'...' if len(product.get('PRODUCT_DESCRIPTION', '')) > 80 else ''}</td>
                                    <td>{product.get('RANGE_LABEL', '')}</td>
                                    <td>{product.get('COMMERCIAL_STATUS', '')}</td>
                                    <td>{product.get('PL_SERVICES', '')}</td>
                                    <td>{product.get('BU_LABEL', '')}</td>
                                </tr>"""
                
                if len(result.products) > 100:
                    html += f"""
                                <tr>
                                    <td colspan="6" style="text-align: center; font-style: italic; background: #f8f9fa;">
                                        <strong>... and {len(result.products) - 100} more products</strong><br>
                                        <small>Total: {len(result.products):,} products found for this document</small>
                                    </td>
                                </tr>"""
                
                html += """
                            </tbody>
                        </table>
                    </div>
                </div>"""
            elif result.success and not result.products:
                html += f"""
                <div id="products-doc-{i}" class="document-products" style="display: none;">
                    <h4>📄 {result.file_name}</h4>
                    <div style="background: #fff3cd; padding: 15px; border-radius: 8px; border-left: 4px solid #ffc107; margin: 15px 0;">
                        <strong>⚠️ No Products Found</strong><br>
                        Extraction method: {result.extraction_method}<br>
                        Ranges detected: {', '.join(result.ranges) if result.ranges else 'None'}<br>
                        Confidence: {result.extraction_confidence:.2f}<br>
                        <small>This may indicate the ranges are not present in the IBcatalogue or need different search criteria.</small>
                    </div>
                </div>"""
            else:
                html += f"""
                <div id="products-doc-{i}" class="document-products" style="display: none;">
                    <h4>📄 {result.file_name}</h4>
                    <div style="background: #f8d7da; padding: 15px; border-radius: 8px; border-left: 4px solid #dc3545; margin: 15px 0;">
                        <strong>❌ Document Processing Failed</strong><br>
                        Error: {result.error}<br>
                        <small>Unable to extract content from this document.</small>
                    </div>
                </div>"""
        
        html += """
            </div>
            
            <div id="analytics" class="tab-content">
                <h3>📈 Enhanced Extraction Analytics</h3>
                <p>Detailed analysis of multi-strategy extraction performance:</p>"""
        
        # Analytics for each document
        for result in results:
            if result.success:
                html += f"""
                <div style="background: #f8f9fa; padding: 15px; border-radius: 8px; margin: 15px 0;">
                    <h4>📄 {result.file_name}</h4>
                    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px;">
                        <div>
                            <strong>🔌 Context Intelligence:</strong><br>
                            Voltage: {result.context.voltage_level or 'Not detected'}<br>
                            Category: {result.context.product_category or 'Not detected'}<br>
                            PL_SERVICES: {result.context.pl_services_hint}<br>
                            Confidence: {result.context.confidence_score:.2f}
                        </div>
                        <div>
                            <strong>🎯 Extraction Performance:</strong><br>
                            Method: {result.extraction_method}<br>
                            Confidence: {result.extraction_confidence:.2f}<br>
                            Ranges Found: {len(result.ranges or [])}<br>
                            Products Matched: {result.product_count:,}
                        </div>
                        <div>
                            <strong>📊 Processing Metrics:</strong><br>
                            Search Reduction: {result.search_space_reduction:.1f}%<br>
                            Processing Time: {result.processing_time_ms:.1f}ms<br>
                            Success Rate: {'100%' if result.success else '0%'}<br>
                            Business Context: {', '.join(result.context.business_context) if result.context.business_context else 'None'}
                        </div>
                    </div>
                    <p><strong>🎯 Extracted Ranges:</strong> {', '.join(result.ranges) if result.ranges else 'None found'}</p>
                </div>"""
        
        html += """
            </div>
        </div>
    </div>
    
    <script>
        let currentDoc = 0;
        
        function showDocument(index) {
            // Update document cards
            document.querySelectorAll('.document-card').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.document-card')[index].classList.add('active');
            currentDoc = index;
            
            // Update products tab to show selected document's products
            showProductsForDocument(index);
        }
        
        function showProductsForDocument(docIndex) {
            // Hide all document product sections
            document.querySelectorAll('.document-products').forEach(el => el.style.display = 'none');
            
            // Show the selected document's products
            const selectedProducts = document.getElementById(`products-doc-${docIndex}`);
            if (selectedProducts) {
                selectedProducts.style.display = 'block';
            }
            
            // Update the products container
            const container = document.getElementById('products-container');
            if (container) {
                container.style.display = 'none';
            }
        }
        
        function showTab(evt, tabName) {
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.tab').forEach(el => el.classList.remove('active'));
            
            document.getElementById(tabName).classList.add('active');
            evt.currentTarget.classList.add('active');
            
            // If switching to products tab, show current document's products
            if (tabName === 'products') {
                showProductsForDocument(currentDoc);
            }
        }
        
        function filterTable(input, tableId) {
            const filter = input.value.toLowerCase();
            const table = document.getElementById(tableId);
            if (!table) return;
            
            const rows = table.getElementsByTagName("tr");
            
            for (let i = 1; i < rows.length; i++) {
                const row = rows[i];
                const cells = row.getElementsByTagName("td");
                let match = false;
                
                for (let j = 0; j < cells.length; j++) {
                    if (cells[j].textContent.toLowerCase().indexOf(filter) > -1) {
                        match = true;
                        break;
                    }
                }
                
                row.style.display = match ? "" : "none";
            }
        }
        
        // Initialize: show first document's products on load
        document.addEventListener('DOMContentLoaded', function() {
            showProductsForDocument(0);
        });
        
        // Keyboard navigation
        document.addEventListener("keydown", function(e) {
            if (e.key === "ArrowUp" && currentDoc > 0) {
                showDocument(currentDoc - 1);
            } else if (e.key === "ArrowDown" && currentDoc < {len(results) - 1}) {
                showDocument(currentDoc + 1);
            }
        });
    </script>
</body>
</html>"""
        
        return html


class SELettersEnhancedPipeline:
    """Enhanced SE Letters Pipeline with Multi-Strategy Range Extraction"""
    
    def __init__(self):
        self.context_analyzer = IntelligentContextAnalyzer()
        self.doc_processor = DocumentProcessor()
        self.db_service = DuckDBService()
        self.html_generator = HTMLReportGenerator()
        self.output_dir = Path("data/output")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def run_pipeline(self, num_docs: int = 5) -> str:
        """Run the enhanced pipeline"""
        print("🚀 SE LETTERS ENHANCED PIPELINE - MULTI-STRATEGY EXTRACTION")
        print("=" * 90)
        print("🧠 Context Intelligence | 🎯 Multi-Strategy Detection | 🚀 DuckDB Ultra-Fast | 📊 Enhanced Reports")
        print()
        
        # Find documents
        docs_dir = Path("data/input/letters")
        doc_files = []
        for pattern in ["*.doc*", "*.pdf"]:
            doc_files.extend(docs_dir.glob(pattern))
            doc_files.extend(docs_dir.glob(f"*/{pattern}"))
            doc_files.extend(docs_dir.glob(f"*/*/{pattern}"))
        
        if not doc_files:
            print("❌ No documents found")
            return ""
        
        # Select random documents
        selected_docs = random.sample(doc_files, min(num_docs, len(doc_files)))
        print(f"📄 Processing {len(selected_docs)} random documents with enhanced extraction")
        
        results = []
        
        for i, doc_file in enumerate(selected_docs, 1):
            print(f"\n🔄 Document {i}/{len(selected_docs)}: {doc_file.name}")
            start_time = time.time()
            
            # 1. Context analysis
            context = self.context_analyzer.analyze_document_context(doc_file)
            print(f"  🧠 Context: {context.voltage_level or 'Unknown'} | {context.product_category or 'Unknown'} | {context.pl_services_hint} (Conf: {context.confidence_score:.2f})")
            
            # 2. Document processing
            doc_result = self.doc_processor.process_document(doc_file)
            
            if not doc_result['success']:
                print(f"  ❌ Failed: {doc_result.get('error', 'Unknown error')}")
                results.append(ProcessingResult(
                    success=False,
                    file_name=doc_file.name,
                    file_path=str(doc_file),
                    file_size=doc_file.stat().st_size,
                    context=context,
                    error=doc_result.get('error', 'Unknown error'),
                    processing_time_ms=(time.time() - start_time) * 1000
                ))
                continue
            
            print(f"  📄 Text: {len(doc_result['content'])} characters")
            
            # 3. Enhanced multi-strategy range extraction
            extraction_result = self.context_analyzer.range_extractor.extract_ranges_comprehensive(
                doc_result['content'], context
            )
            
            ranges = extraction_result['ranges']
            extraction_method = extraction_result['extraction_method']
            extraction_confidence = extraction_result['extraction_confidence']
            
            print(f"  🎯 Multi-strategy extraction: {ranges} | Method: {extraction_method} | Conf: {extraction_confidence:.2f}")
            
            # 4. Product search with intelligence
            search_reduction = self.db_service.calculate_search_space_reduction(context)
            products = self.db_service.find_products_with_context(ranges, context)
            processing_time = time.time() - start_time
            
            print(f"  🚀 DuckDB: {len(products)} products | {search_reduction:.1f}% reduction | {processing_time*1000:.1f}ms")
            
            # Create result
            result = ProcessingResult(
                success=True,
                file_name=doc_file.name,
                file_path=str(doc_file),
                file_size=doc_file.stat().st_size,
                context=context,
                content=doc_result['content'],
                ranges=ranges,
                products=products,
                product_count=len(products),
                processing_time_ms=processing_time * 1000,
                search_space_reduction=search_reduction,
                extraction_method=extraction_method,
                extraction_confidence=extraction_confidence
            )
            
            results.append(result)
        
        # Generate HTML report
        print(f"\n📊 Generating enhanced HTML report...")
        html_content = self.html_generator.generate_report(results)
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = self.output_dir / f"SE_Letters_Enhanced_Pipeline_Report_{timestamp}.html"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Summary
        successful = len([r for r in results if r.success])
        total_products = sum(r.product_count for r in results if r.success)
        total_ranges = sum(len(r.ranges or []) for r in results if r.success)
        avg_reduction = sum(r.search_space_reduction for r in results if r.success) / max(successful, 1)
        avg_confidence = sum(r.extraction_confidence for r in results if r.success) / max(successful, 1)
        
        print(f"\n🏆 ENHANCED PIPELINE COMPLETE")
        print(f"📊 Documents: {len(results)} ({successful} successful)")
        print(f"🎯 Ranges extracted: {total_ranges}")
        print(f"🛠️ Products found: {total_products:,}")
        print(f"📉 Average search reduction: {avg_reduction:.1f}%")
        print(f"🎯 Average extraction confidence: {avg_confidence:.2f}")
        print(f"📁 Report: {report_path}")
        
        return str(report_path)
    
    def close(self):
        """Close connections"""
        self.db_service.close()


if __name__ == "__main__":
    pipeline = SELettersEnhancedPipeline()
    try:
        report_path = pipeline.run_pipeline(5)
        if report_path:
            print(f"\n🌐 Opening enhanced report in browser...")
            import subprocess
            subprocess.run(["open", report_path])
    finally:
        pipeline.close() 